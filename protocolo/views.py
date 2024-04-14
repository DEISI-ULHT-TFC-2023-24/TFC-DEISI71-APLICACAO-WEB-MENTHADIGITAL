from django.contrib.auth import logout
import io
from platform import python_compiler
from django.http import HttpResponseRedirect, JsonResponse, FileResponse
from django.shortcuts import render, HttpResponse, redirect
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login
from matplotlib import pyplot as plt
from .models import Protocol, Part, Area, Instrument, Dimension, Section, Question, Resolution, Answer, PossibleAnswer, Risk
from django.contrib.auth.models import User
from django.contrib.auth.models import Group as DjangoGroup
from django.urls import reverse
from .functions import *
from .forms import *
from diario.models import *
import json
import os


# Other Imports
import plotly.graph_objects as go
import plotly
import pandas as pd
import time
import hashlib
import random
from PIL import Image, ImageDraw, ImageFont
from docx.shared import Pt


from bs4 import BeautifulSoup

import numpy as np
import plotly.offline as py_offline


# word Imports to PDF
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


import os


# Create your views here.
def remove_empty_keys_from_post(post_data):
    return {key: value for key, value in post_data.items() if value != ''}


def inserir():
    a1 = PossibleAnswer.objects.get(pk=310)
    a2 = PossibleAnswer.objects.get(pk=311)
    a3 = PossibleAnswer.objects.get(pk=312)
    a4 = PossibleAnswer.objects.get(pk=313)
    a5 = PossibleAnswer.objects.get(pk=314)
    a6 = PossibleAnswer.objects.get(pk=315)

    qorder = 1
    qtype = 3
    section = Section.objects.get(pk=116)
    text_list = [
        "De quantos familiares se sente próximo de tal forma que possa ligar-lhes para pedir ajuda?",
        "Com quantos familiares se sente à vontade para falar sobre assuntos pessoais?",
        "Quantos amigos vê ou fala pelo menos uma vez por mês?",
        "De quantos amigos se sente próximo de tal forma que possa ligar-lhes a pedir ajuda?",
        "Com quantos amigos se sente à vontade para falar sobre assuntos pessoais?"
    ]

    for text in text_list:
        qorder += 1
        q = Question(
            name=text,
            order=qorder,
            question_type=qtype,
            section=section,
            quotation_max=5,
        )
        q.save()
        q.possible_answers.add(a1)
        q.possible_answers.add(a2)
        q.possible_answers.add(a3)
        q.possible_answers.add(a4)
        q.possible_answers.add(a5)
        q.possible_answers.add(a6)
        q.save()


def verificar_form_risk(value):
    if value is None or len(value) == 0:
        if type(value) is str:
            return ""
        else:
            return 0
    if type(value) is str:
        value = value.replace(".", ",")
    return value


def has_cuidador_flag(flag):
    return flag == 'cuidador'


def get_participant_or_cuidador(flag_boolean, id):
    if flag_boolean:
        return Cuidador.objects.get(pk=id)

    return Participante.objects.get(pk=id)


@login_required(login_url='login')
def dashboard_view(request):
    doctor = request.user
    participants = Participante.objects.filter(avaliador=doctor)
    resolutions = Resolution.objects.filter(doctor=doctor)

    context = {'participants': participants, 'resolutions': resolutions}
    return render(request, 'protocolo/dashboard.html', context)


@login_required(login_url='login')
def dashboard_content_view(request):
    doctor = request.user
    participants = Participante.objects.filter(avaliador=doctor)
    resolutions = Resolution.objects.filter(doctor=doctor)

    context = {'participants': participants, 'resolutions': resolutions}
    return render(request, 'protocolo/dashboardcontent.html', context)


@login_required(login_url='login')
def teste(resquest):
    return render(resquest, 'protocolo/teste.html')


@login_required(login_url='login')
def protocolos_view(request):
    context = {'protocolos': Protocol.objects.all().order_by('order')}
    return render(request, 'protocolo/protocolos.html', context)


@login_required(login_url='login')
def menu_protocolo_view(request):
    return render(request, 'mentha/app-list.html')


@login_required(login_url='login')
def popup_view(request):
    return render(request, 'protocolo/popupparts.html')


@login_required(login_url='login')
def registo_view(request):
    doctor = request.user

    if request.method == 'POST':
        try:
            user = Avaliador.objects.get(user=doctor)
        except Avaliador.DoesNotExist:
            try:
                user = Administrador.objects.get(user=doctor)
            except Administrador.DoesNotExist:
                user = None

        referenciacao = user.reference
        new_user = User()
        new_user.username = request.POST.get('username').replace(" ", "")
        new_user.password = request.POST.get('password')
        new_user.email = request.POST.get('email')
        new_user.save()

        my_group = DjangoGroup.objects.get(name='Participante')
        my_group.user_set.add(new_user)

        new_registo = InformacaoSensivel()
        new_registo.nome = request.POST.get('nome')
        new_registo.email = request.POST.get('email')
        new_registo.telemovel = request.POST.get('telemovel')
        new_registo.save()

        new_participante = Participante()
        new_participante.user = new_user
        new_participante.sexo = request.POST.get('sexo')
        new_participante.info_sensivel = new_registo
        new_participante.referenciacao = referenciacao
        new_participante.avaliador = doctor
        nascimento_str = request.POST.get('nascimento')
        nascimento_datetime = datetime.strptime(nascimento_str, '%Y-%m-%d')
        new_participante.nascimento = nascimento_datetime
        new_participante.save()

        return redirect('protocolo:participants')

    context = {}
    return render(request, 'protocolo/participantes_registo.html')


@login_required(login_url='login')
def avaliadores_registo_view(request):
    if request.method == 'POST':
        admin = Administrador.objects.get(user=request.user)

        new_user = User()

        new_user.username = request.POST.get('username')
        new_user.password = request.POST.get('password')
        new_user.email = request.POST.get('email')

        new_user.first_name = request.POST.get('primeiro_nome')
        new_user.last_name = request.POST.get('ultimo_nome')
        new_user.save()

        new_avaliador = Avaliador()
        new_avaliador.user = new_user
        new_avaliador.reference = admin.reference
        new_avaliador.save()

        my_group = DjangoGroup.objects.get(name='Avaliador')
        my_group.user_set.add(new_user)

        return redirect('protocolo:participants')

    context = {}
    return render(request, 'protocolo/avaliadores_registo.html')


@login_required(login_url='login')
def incrementar(request, part_id, participant_id):
    part = Part.objects.get(pk=part_id)
    particitant = Participante.objects.get(pk=participant_id)
    nome = part.nome + datetime.now
    partParticipant = ParteDoUtilizador(
        part=part, participante=particitant, nome=nome)

    return render(request, 'protocolo/parts.html')


@login_required(login_url='login')
def participant_risk_view(request, protocol_id):
    doctor = request.user
    protocolo = Protocol.objects.get(pk=protocol_id)
    participants = Participante.objects.filter(avaliador=doctor)
    resolutions = Resolution.objects.filter(doctor=doctor)

    context = {'participants': participants,
               'resolutions': resolutions, 'protocolo': protocolo}
    return render(request, 'protocolo/protocolo_participants_risk.html', context)
# funcao view para adicionar uma parte a um participante


@login_required(login_url='login')
def nova_pagina_risk_report(request):
    form_report_risk = Risk.objects.all()
    return render(request, 'protocolo/visualizacao_risk.html', {'form_report_risk': form_report_risk})


@login_required(login_url='login')
def parte_do_utilizador_add_view(request):
    if request.method == 'POST':
        partId = request.POST.get('partId')
        part = Part.objects.get(pk=partId)
        patient = Participante.objects.get(pk=request.POST.get('patientId'))
        parteDoUtilizador = ParteDoUtilizador(part=part, participante=patient)
        parteDoUtilizador.save()

    context = {
        'parteDoUtilizador': parteDoUtilizador,
        'part': part,

    }
    return HttpResponse('200')
    # return JsonResponse({}, status=201)


@login_required(login_url='login')
def parte_view(request):
    parte = Part.objects.all()  # participante = patient)
    # patient = Participante.objects.get(pk=patient_id)
    botao_clicado = int(request.GET.get('botao_clicado ', 0))
    context = {
        'parte': parte,
        'botao_clicado': botao_clicado,
        # 'patient':patient
    }
    return render(request, 'protocolo/parts.html', context)


@login_required(login_url='login')
def parts_view(request, protocol_id, patient_id, is_cuidador='is_participante'):
    start = time.time()
    protocol = Protocol.objects.get(pk=protocol_id)
    if is_cuidador == 'is_cuidador':
        cuidador = Cuidador.objects.get(pk=patient_id)
        # Mudar request.user para o patient depois
        resolutions = Resolution.objects.filter(
            doctor=request.user, cuidador=cuidador)
        parts = ParteDoUtilizador.objects.filter(cuidador=cuidador)
        patient = cuidador  # talvez precise disto para o url renderizar
    else:
        patient = Participante.objects.get(pk=patient_id)
        # Mudar request.user para o patient depois
        resolutions = Resolution.objects.filter(
            doctor=request.user, patient=patient)
        parts = ParteDoUtilizador.objects.filter(participante=patient)

    parte = Part.objects.all()
    risk_area = Area.objects.get(id=47)
    pergunta_risk = Question.objects.get(id=189)
    print(pergunta_risk)
    # statistics
    answered_list = []
    percentage_list = []
    # data = parte_parte.get().data
    for parteDoUtilizador in parts:
        resolution = resolutions.filter(part=parteDoUtilizador)
        if not resolution:
            answered_list.append(0)
            percentage_list.append(0)
        else:
            s = resolution.get().statistics
            answered_list.append(s.get('total_answered'))
            percentage_list.append(s.get('total_percentage'))

    end = time.time()
    print("Parts", (end - start))
    context = {'parts': zip(parts, answered_list, percentage_list),
               'parte': parte,
               'risk_area': risk_area,
               'protocol': protocol, 'resolutions': resolutions,
               'patient': patient,
               'pergunta_risk': pergunta_risk,
               'is_cuidador': is_cuidador,
               }
    return render(request, 'protocolo/parts.html', context)


@login_required(login_url='login')
def areas_view(request, protocol_id, part_id, patient_id, is_cuidador='is_participante'):
    start = time.time()
    protocol = Protocol.objects.get(pk=protocol_id)
    parteDoUtilizador = ParteDoUtilizador.objects.get(pk=part_id)
    part = parteDoUtilizador.part
    areas = Area.objects.filter(part=part).order_by('order')
    caderno_estimulos = part.caderno_estimulos

    if is_cuidador == 'is_cuidador':
        cuidador = Cuidador.objects.get(pk=patient_id)
        patient = cuidador
        r = Resolution.objects.filter(
            cuidador=cuidador, part=parteDoUtilizador).get()
    else:
        patient = Participante.objects.get(pk=patient_id)
        r = Resolution.objects.filter(
            patient=patient, part=parteDoUtilizador).get()

    rel_q = Question.objects.filter(name="Relação com o Avaliador").get()
    coop_q = Question.objects.filter(
        name="Cooperação dada na entrevista").get()
    qs_q = Question.objects.filter(name="Questionário Sociodemográfico").get()

    answered_list = []
    percentage_list = []
    s = r.statistics
    for area in areas:
        if s.get(f'{area.id}') is not None:
            answered_list.append(s.get(f'{area.id}').get('answered'))
            percentage_list.append(s.get(f'{area.id}').get('percentage'))
    end = time.time()
    print("Parts", (end - start))
    context = {'areas': zip(areas, answered_list, percentage_list),
               'part': parteDoUtilizador,
               'protocol': protocol,
               'resolution': r.id,
               'patient': patient,
               'coop': coop_q,
               'rel': rel_q,
               'qs': qs_q,
               'is_cuidador': is_cuidador,
               'caderno_estimulos': caderno_estimulos}
    return render(request, 'protocolo/areas.html', context)


@login_required(login_url='login')
def instruments_view(request, protocol_id, part_id, area_id, patient_id, is_cuidador='is_participante'):
    start = time.time()
    protocol = Protocol.objects.get(pk=protocol_id)
    parteDoUtilizador = ParteDoUtilizador.objects.get(pk=part_id)
    area = Area.objects.get(pk=area_id)

    if is_cuidador == 'is_cuidador':
        cuidador = Cuidador.objects.get(pk=patient_id)
        r = Resolution.objects.get(cuidador=cuidador, part=parteDoUtilizador)
        patient = cuidador
    else:
        patient = Participante.objects.get(pk=patient_id)
        r = Resolution.objects.get(patient=patient, part=parteDoUtilizador)

    instruments = Instrument.objects.filter(area=area_id).order_by('order')

    answered_list = []
    percentage_list = []
    quotation_list = []
    s = r.statistics
    for instrument in instruments:
        if s.get(f'{area.id}') is not None:
            answered_list.append(s.get(f'{area.id}').get(
                f'{instrument.id}').get('answered'))
            percentage_list.append(s.get(f'{area.id}').get(
                f'{instrument.id}').get('percentage'))
            quotation_list.append(s.get(f'{area.id}').get(
                f'{instrument.id}').get('quotation'))
    # print(answered_list)
    # print(percentage_list)
    end = time.time()
    print("Instruments", (end - start))

    context = {'area': area, 'part': parteDoUtilizador, 'protocol': protocol, 'protocol_id': protocol_id, 'part_id': part_id,
               'area_id': area_id, 'patient_id': patient_id,
               'instruments': [e for e in zip(instruments, answered_list, percentage_list, quotation_list)],
               'resolution': r.id,
               'patient': patient, 'is_cuidador': is_cuidador}

    renderizado = render(request, 'protocolo/instruments.html', context)
    end = time.time()
    print("Instruments", (end - start))
    return renderizado


@login_required(login_url='login')
def dimensions_view(request, protocol_id, part_id, area_id, instrument_id, patient_id, is_cuidador='is_participante'):
    start = time.time()
    protocol = Protocol.objects.get(pk=protocol_id)
    parteDoUtilizador = ParteDoUtilizador.objects.get(pk=part_id)
    area = Area.objects.get(pk=area_id)
    instrument = Instrument.objects.get(pk=instrument_id)
    dimensions = Dimension.objects.filter(
        instrument=instrument_id).order_by('order')

    if is_cuidador == 'is_cuidador':
        cuidador = Cuidador.objects.get(pk=patient_id)
        patient = cuidador
        r = Resolution.objects.get(cuidador=cuidador, part=parteDoUtilizador)
    else:
        patient = Participante.objects.get(pk=patient_id)
        r = Resolution.objects.get(patient=patient, part=parteDoUtilizador)

    if len(dimensions) == 1:
        return redirect('protocolo:sections', protocol_id, part_id, area_id, instrument_id, dimensions.get().id, patient_id, is_cuidador)

    answered_list = []
    percentage_list = []
    quotation_list = []
    s = r.statistics
    for dimension in dimensions:
        if s.get(f'{area.id}') is not None:
            answered_list.append(s.get(f'{area.id}').get(
                f'{instrument.id}').get(f'{dimension.id}').get('answered'))
            percentage_list.append(s.get(f'{area.id}').get(
                f'{instrument.id}').get(f'{dimension.id}').get('percentage'))
            quotation_list.append(s.get(f'{area.id}').get(
                f'{instrument.id}').get(f'{dimension.id}').get('quotation'))
    # print(answered_list)
    # print(percentage_list)
    end = time.time()
    print("Dimensions", (end - start))
    context = {'area': area, 'part': parteDoUtilizador, 'protocol': protocol, 'instrument': instrument,
               'dimensions': zip(dimensions, answered_list, percentage_list, quotation_list), 'resolution': r.id,
               'patient': patient, 'is_cuidador': is_cuidador}
    return render(request, 'protocolo/dimensions.html', context)


@login_required(login_url='login')
def sections_view(request, protocol_id, part_id, area_id, instrument_id, dimension_id, patient_id, is_cuidador='is_participante'):
    start = time.time()
    protocol = Protocol.objects.get(pk=protocol_id)
    parteDoUtilizador = ParteDoUtilizador.objects.get(pk=part_id)
    area = Area.objects.get(pk=area_id)
    instrument = Instrument.objects.get(pk=instrument_id)
    dimension = Dimension.objects.get(pk=dimension_id)

    if is_cuidador == 'is_cuidador':
        cuidador = Cuidador.objects.get(pk=patient_id)
        patient = cuidador
        r = Resolution.objects.get(cuidador=patient, part=parteDoUtilizador)
    else:
        patient = Participante.objects.get(pk=patient_id)
        r = Resolution.objects.get(patient=patient, part=parteDoUtilizador)

    sections = Section.objects.filter(dimension=dimension_id).order_by('order')
    if len(sections) == 1:
        return redirect('protocolo:question', protocol_id, part_id, area_id, instrument_id, dimension_id, sections.get().id, patient_id, is_cuidador)

    answered_list = []
    percentage_list = []
    quotation_list = []
    s = r.statistics
    for section in sections:
        if s.get(f'{area.id}') is not None:
            answered_list.append(
                s.get(f'{area.id}').get(f'{instrument.id}').get(f'{dimension.id}').get(f'{section.id}').get('answered'))
            percentage_list.append(
                s.get(f'{area.id}').get(f'{instrument.id}').get(f'{dimension.id}').get(f'{section.id}').get(
                    'percentage'))
            quotation_list.append(
                s.get(f'{area.id}').get(f'{instrument.id}').get(f'{dimension.id}').get(f'{section.id}').get(
                    'quotation'))

    if len(sections) == 1:
        return redirect(question_view)

    end = time.time()
    print("Sections", (end - start))
    context = {'area': area, 'part': parteDoUtilizador, 'protocol': protocol, 'instrument': instrument, 'dimension': dimension,
               'sections': zip(sections, answered_list, percentage_list, quotation_list), 'resolution': r.id,
               'patient': patient, 'is_cuidador': is_cuidador}

    # print(instrument.number_of_dimensions)

    return render(request, 'protocolo/sections.html', context)


@login_required(login_url='login')
def question_view(request, protocol_id, part_id, area_id, instrument_id, dimension_id, section_id, patient_id, is_cuidador='is_participante'):
    start = time.time()
    doencas = Doenca.objects.all()
    protocol = Protocol.objects.get(pk=protocol_id)
    parteDoUtilizador = ParteDoUtilizador.objects.get(pk=part_id)
    area = Area.objects.get(pk=area_id)
    instrument = Instrument.objects.get(pk=instrument_id)
    dimension = Dimension.objects.get(pk=dimension_id)
    section = Section.objects.get(pk=section_id)
    question = Question.objects.filter(section=section).first()
    form = uploadAnswerForm(request.POST or None)
    form_risk = FormRisk(request.POST or None)

    if is_cuidador == 'is_cuidador':
        cuidador = Cuidador.objects.get(pk=patient_id)
        patient = cuidador
        r = Resolution.objects.get(cuidador=cuidador, part=parteDoUtilizador)
    else:
        patient = Participante.objects.get(pk=patient_id)
        r = Resolution.objects.get(patient=patient, part=parteDoUtilizador)

    boolean_pressao = False
    boolean_idade = False

    if question.question_type == 11:
        return redirect('protocolo:participante_update', patient_id)

    if question.question_type == 13:
        return redirect('protocolo:cuidador_update', patient_id)

    answers = Answer.objects.filter(resolution=r)

    context = {'area': area, 'part': parteDoUtilizador, 'protocol': protocol, 'instrument': instrument, 'dimension': dimension,
               'section': section, 'question': question, 'form': form, 'resolution': r.id, 'answers': answers,
               'patient': patient, 'doencas': doencas, 'form_risk': form_risk,
               'is_cuidador': is_cuidador,
               }

    if question.question_type == 10:
        if len(answers.filter(question=question)) > 0:
            a = answers.filter(question=question).get()
            json_answer = json.loads(a.text_answer)
            context['Sequenciacao'] = json_answer.get('Sequenciacao')
            context['Perseverativos'] = json_answer.get('Perseverativos')
            context['Proximidade'] = json_answer.get('Proximidade')
            context['Tempo'] = json_answer.get('Tempo')

    if question.question_type == 3:
        question_list = []
        answered_ids = []
        for question in Question.objects.filter(section=section.id):
            question_list.append(question)
            for answer in answers:
                if question == answer.question:
                    answered_ids.append(question.id)

        # Esta parte permite dividir o question type 3 em dois
        # Um que tem sempre as mesmas respostas, e mostrará a página como uma tabela
        # Outro que as respostas são diferentes e mostrará varias perguntas individualmente, todas na mesma página
        ans = 0
        equal = True
        qset = []
        for question in Question.objects.filter(section=section.id):
            if ans == 0:
                ans = 1
                qset = question.possible_answer_name_list  # print(qset)
            else:
                # print(f"{question.possible_answer_name_list} vs {qset}")
                equal = question.possible_answer_name_list == qset

        context['equal_answers'] = equal
        context['question_list'] = question_list

    existing_answer_id = []
    for answer in answers:
        if answer.resolution == r:
            if answer.question == question:
                if answer.multiple_choice_answer is not None:
                    existing_answer_id.append(answer.multiple_choice_answer.id)
                    context['existing_answer_id'] = existing_answer_id
                if answer.text_answer is not None:
                    form.initial.update({'text_answer': answer.text_answer})
                if answer.submitted_answer:
                    context['submitted_answer'] = answer.submitted_answer.url
                if answer.notes is not None:
                    context['notes'] = answer.notes
                if answer.quotation is not None:
                    context['quotation'] = answer.quotation
                if len(answer.MCCAnswer.all()) >= 1:
                    list = []
                    for mca in answer.MCCAnswer.all():
                        list.append(mca.choice.id)
                    existing_answer_id = list
                    context['existing_answer_id'] = list
    existing_risk = None
    for risk in Risk.objects.all():
        if risk.parteDoUtilizador == parteDoUtilizador:
            existing_risk = risk
            context['boolean_idade'] = existing_risk.idade >= 40 and existing_risk.idade <= 89
            context['boolean_pressao'] = existing_risk.pressao_arterial >= 100 and existing_risk.pressao_arterial <= 179
    context['existing_risk'] = existing_risk

    if request.method == 'POST':
        existing_answer = None

        for answer in answers:
            if answer.question == question:
                existing_answer = answer

        if question.question_type == 1 or question.question_type == 9:
            id_answer = request.POST.get("choice")
            # r = Resolution.objects.get(part=parteDoUtilizador, patient=patient)
            if existing_answer is None:
                # cria uma nova associação
                new_answer = Answer(question=question, multiple_choice_answer=PossibleAnswer.objects.get(pk=id_answer),
                                    resolution=r)
                quotation = new_answer.multiple_choice_answer.quotation
                new_answer.quotation = quotation
                new_answer.notes = request.POST.get('notes')
                new_answer.save()
                r.increment_statistics(f'{part_id}', f'{area_id}', f'{instrument_id}', f'{dimension_id}',
                                       f'{section_id}')
                r.change_quotation(
                    f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}', quotation)
            else:
                # modifica a associação existente
                existing_answer.multiple_choice_answer = PossibleAnswer.objects.get(
                    pk=id_answer)
                quotation = existing_answer.multiple_choice_answer.quotation
                existing_answer.quotation = quotation
                existing_answer.notes = request.POST.get('notes')
                existing_answer.save()
                r.change_quotation(
                    f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}', quotation)

        elif question.question_type == 2:
            form = uploadAnswerForm(request.POST, files=request.FILES)
            if form.is_valid():
                new_answer = Answer()
                new_answer.submitted_answer = form.cleaned_data['submitted_answer']
                new_answer.text_answer = form.cleaned_data['text_answer']
                new_answer.quotation = form.cleaned_data['quotation']
                new_answer.notes = form.cleaned_data['notes']
                new_answer.question = question
                new_answer.resolution = r

                if existing_answer is None:
                    # cria uma nova resposta
                    r.increment_statistics(f'{part_id}', f'{area_id}', f'{instrument_id}', f'{dimension_id}',
                                           f'{section_id}')
                    new_answer.resolution = r
                    new_answer.save()
                else:
                    # modifica a resposta existente
                    existing_answer.submitted_answer = new_answer.submitted_answer
                    existing_answer.text_answer = new_answer.text_answer
                    existing_answer.quotation = new_answer.quotation
                    existing_answer.notes = new_answer.notes
                    existing_answer.save()

                r.change_quotation(f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}',
                                   new_answer.quotation)

        elif question.question_type == 3:
            for key in request.POST:
                if 'choice' in str(key):
                    k, question_id = key.split('-')
                    q = Question.objects.get(pk=question_id)
                    existing_answers_list = []

                    if q.id in answered_ids:
                        a = Answer.objects.filter(
                            resolution=r, question=q).get()
                        # modifica a associação existente
                        a.multiple_choice_answer = PossibleAnswer.objects.get(
                            pk=request.POST.get(key))
                        quotation = a.multiple_choice_answer.quotation
                        a.quotation = quotation
                        a.notes = request.POST.get('notes')
                        a.save()
                        r.change_quotation(f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}',
                                           quotation)
                    else:
                        a = Answer()
                        a.resolution = r
                        a.multiple_choice_answer = PossibleAnswer.objects.get(
                            pk=request.POST.get(key))
                        a.question = q
                        quotation = a.multiple_choice_answer.quotation
                        a.quotation = quotation
                        a.notes = request.POST.get('notes')
                        a.save()
                        r.increment_statistics(f'{part_id}', f'{area_id}', f'{instrument_id}', f'{dimension_id}',
                                               f'{section_id}')
                        r.change_quotation(f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}',
                                           quotation)

        elif question.question_type == 4 or question.question_type == 6 or question.question_type == 7 or question.question_type == 8:
            existing = False
            if len(Answer.objects.filter(question=question, resolution=r)) >= 1:
                a = Answer.objects.filter(
                    question=question, resolution=r).get()
                a.delete()
                existing = True
            a = Answer()
            a.resolution = r
            a.question = question
            a.notes = request.POST.get('notes')
            q = 0
            a.save()
            for id in request.POST.getlist("choice"):
                c = MultipleChoicesCheckbox()
                c.answer = a
                pa = PossibleAnswer.objects.filter(pk=id).get()
                c.choice = pa
                q = q + c.choice.quotation
                c.save()

            if 'Repetição I' in question.name:
                l = len(request.POST.getlist("choice"))
                if l == 4:
                    q = 2
                elif l == 3:
                    q = 1
                else:
                    q = 0

            a.quotation = q
            a.save()

            if existing:
                r.change_quotation(
                    f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}', q)
            else:
                r.increment_statistics(f'{part_id}', f'{area_id}', f'{instrument_id}', f'{dimension_id}',
                                       f'{section_id}')
                r.change_quotation(
                    f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}', q)

        elif question.question_type == 5:
            # print(request.POST)
            existing = False
            if len(Answer.objects.filter(question=question, resolution=r)) >= 1:
                a = Answer.objects.filter(
                    question=question, resolution=r).get()
                a.delete()
                existing = True
            a = Answer()
            a.resolution = r
            a.question = question
            a.notes = request.POST.get('notes')
            a.save()
            counter = 0
            for i in range(1, 5):
                text_area = request.POST.get(str(i))
                if len(text_area) > 1:
                    counter += len(text_area.split(","))
                    # print(text_area)
                    # print(counter)
                    ti = TextInputAnswer()
                    ti.answer = a
                    ti.seconds = i
                    ti.text = text_area
                    ti.save()
            q = calculate_timer_quotation(question, counter)
            a.quotation = q
            a.save()

            if existing:
                r.change_quotation(
                    f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}', q)
            else:
                r.increment_statistics(f'{part_id}', f'{area_id}', f'{instrument_id}', f'{dimension_id}',
                                       f'{section_id}')
                r.change_quotation(
                    f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}', q)

        elif question.question_type == 10:
            erro_sequenciacao = request.POST.get('sequenciacao')
            erro_perseverativos = request.POST.get('perseverativos')
            erro_proximidade = request.POST.get('proximidade')
            tempo = request.POST.get('tempo')
            # print(erro_proximidade, erro_sequenciacao, erro_perseverativos, tempo)
            json_answer = {'Sequenciacao': erro_sequenciacao,
                           'Perseverativos': erro_perseverativos,
                           'Proximidade': erro_proximidade,
                           'Tempo': tempo
                           }

            new_answer = Answer()
            new_answer.question = question
            new_answer.resolution = r
            new_answer.text_answer = json.dumps(json_answer)
            new_answer.quotation = int(tempo) - (
                int(erro_sequenciacao) + int(erro_perseverativos) + int(erro_proximidade))

            if existing_answer is None:
                # cria uma nova resposta
                r.increment_statistics(f'{part_id}', f'{area_id}', f'{instrument_id}', f'{dimension_id}',
                                       f'{section_id}')
                new_answer.save()
            else:
                # modifica a resposta existente
                existing_answer.submitted_answer = new_answer.submitted_answer
                existing_answer.text_answer = new_answer.text_answer
                existing_answer.quotation = new_answer.quotation
                existing_answer.notes = new_answer.notes
                existing_answer.save()

            r.change_quotation(f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}',
                               new_answer.quotation)

        elif question.question_type == 11:
            if request.POST.get('sexo'):
                patient.sexo = request.POST.get('sexo')
            if request.POST.get('nacionalidade'):
                patient.nacionalidade = request.POST.get('nacionalidade')
            if request.POST.get('nascimento'):
                patient.nascimento = request.POST.get('nascimento')
            if request.POST.get('escolaridade'):
                patient.escolaridade = request.POST.get('escolaridade')
            if request.POST.get('residencia'):
                patient.residencia = request.POST.get('residencia')
            if request.POST.get('laboral'):
                patient.situacaoLaboral = request.POST.get('laboral')
            if request.POST.get('profissao'):
                patient.profissaoPrincipal = request.POST.get('profissao')
            if request.POST.get('diagnosticoPrincipal'):
                patient.diagnosticoPrincipal = request.POST.get('diagnosticoPrincipal')
            if request.POST.get('comorbilidades'):
                patient.comorbilidades = request.POST.get('comorbilidades')
            if request.POST.get('economina'):
                patient.situacaoEconomica = request.POST.get('economica')
            if request.POST.get('civil'):
                patient.estadoCivil = request.POST.get('civil')
            if request.POST.get('agregado'):
                patient.agregadoFamiliar = request.POST.get('agregado')
            if request.POST.get('temFilho'):
                patient.temFilhos = request.POST.get('filhos') == 'sim'
            if request.POST.get('nrFilhos'):
                patient.nrFilhos = int(request.POST.get('nrFilhos'))
            if request.POST.get('saude'):
                patient.autoAvaliacaoEstadoSaude = request.POST.get('saude')
            if request.POST.get('dadosCuidador'):
                patient.dadosCuidador = request.POST.get('dadosCuidador')
            if request.POST.get('doenca') is not None:
                patient.diagnosticos.clear()
                for id in request.POST.get('doenca'):
                    patient.diagnosticos.add(Doenca.objects.get(id=id))

            patient.save()
            r.increment_statistics(f'{part_id}', f'{area_id}', f'{instrument_id}', f'{dimension_id}',
                                   f'{section_id}')

        elif question.question_type == 12:
            concluido_risk = True
            ris = None
            guardar_genero = ""
            cwd = os.getcwd()
            cwd2 = os.path.join(cwd, 'protocolo', 'static',
                                'protocolo', 'data_risk')
            file_path_men = os.path.join(cwd2, 'risk_men.json')
            file_path_women = os.path.join(cwd2, 'risk_women.json')
            new_risk = Risk()
            new_risk.idade = verificar_form_risk(request.POST.get('idade'))
            if eval(new_risk.idade) >= 40 and eval(new_risk.idade) <= 89:
                boolean_idade = True
            else:
                boolean_idade = False
            new_risk.sexo = verificar_form_risk(request.POST.get('sexo'))
            new_risk.peso = float(request.POST.get('peso'))
            new_risk.altura = verificar_form_risk(request.POST.get('altura'))
            
            if request.POST.get('altura') == '':
                new_risk.altura = 0
            else:
                new_risk.altura = int(new_risk.altura)
            new_risk.horas_jejum = verificar_form_risk(request.POST.get(
                'horas_jejum'))
            new_risk.pressao_arterial = verificar_form_risk(
                request.POST.get('pressao_arterial'))
            new_risk.colestrol_total = verificar_form_risk(
                request.POST.get('colestrol_total'))
            if request.POST.get('colestrol_total') == '':
                new_risk.colestrol_total = 0
            new_risk.colestrol_total = int(new_risk.colestrol_total)
            if eval(new_risk.pressao_arterial) <= 179 and eval(new_risk.pressao_arterial) >= 100:
                boolean_pressao = True
            else:
                boolean_pressao = False
            new_risk.tg = verificar_form_risk(request.POST.get('tg'))
            if request.POST.get('tg') == '':
                new_risk.tg = 0
            new_risk.tg = int(new_risk.tg)
            new_risk.ldl = verificar_form_risk(request.POST.get('ldl'))
            if request.POST.get('ldl') == '':
                new_risk.ldl = 0
            new_risk.ldl = int(new_risk.ldl)
            new_risk.cholhdl = float(request.POST.get('cholhdl'))
            print(request.POST.get('cholhdl'))
            print(new_risk.cholhdl)
            if request.POST.get('cholhdl') == '':
                new_risk.cholhdl = 0
            # testes dos novos parâmetros eag
            new_risk.ifcc_hba1 = float(request.POST.get('ifcc_hba1'))
            if request.POST.get('ifcc_hba1') == '':
                new_risk.ifcc_hba1 = 0
            new_risk.ngsp_hba1 = float(request.POST.get('ngsp_hba1'))
            if request.POST.get('ngsp_hba1') == '':
                new_risk.ngsp_hba1 = 0
            new_risk.eag_hba1 = float(request.POST.get('eag_hba1'))
            if request.POST.get('eag_hba1') == '':
                new_risk.eag_hba1 = 0
        
            # Save the instance to the database
            new_risk.save()
            # testes dos novos parâmetros ngsp
            print(request.POST.get('ngsp_hba1'))
            print(new_risk.ngsp_hba1)
            if request.POST.get('ngsp_hba1') == '':
                new_risk.ngsp_hba1 = 0
            
            new_risk.batimentos = verificar_form_risk(
                request.POST.get('batimentos'))
            if request.POST.get('batimentos') == '':
                new_risk.batimentos = 0
            new_risk.colestrol_hdl = 0 if verificar_form_risk(request.POST.get(
                'colestrol_hdl')) == '' else verificar_form_risk(request.POST.get('colestrol_hdl'))
            new_risk.colestrol_nao_hdl = verificar_form_risk(
                request.POST.get('colestrol_nao_hdl'))
            if request.POST.get('colestrol_nao_hdl') == '':
                new_risk.colestrol_nao_hdl = 1
            new_risk.colestrol_nao_hdl = int(new_risk.colestrol_nao_hdl)
            new_risk.fumador = verificar_form_risk(request.POST.get('fumador'))
            new_risk.diabetes = verificar_form_risk(
                request.POST.get('diabetes'))
            new_risk.pat_id = verificar_form_risk(request.POST.get('pat'))
            if request.POST.get('pat') == '':
                new_risk.pat_id = 'PID00000'
            anos_diabete = verificar_form_risk(
                request.POST.get('anos_diabetes'))
            if anos_diabete == '':
                anos_diabete = 0
            new_risk.anos_diabetes = anos_diabete
            new_risk.avc = verificar_form_risk(request.POST.get('avc'))
            new_risk.enfarte = verificar_form_risk(request.POST.get('enfarte'))
            new_risk.pre_diabetico = verificar_form_risk(
                request.POST.get('pre_diabetico'))
            new_risk.pergunta_cardiovascular = verificar_form_risk(
                request.POST.get('pergunta_cardiovascular'))
            new_risk.doenca_rins = verificar_form_risk(
                request.POST.get('doenca_rins'))
            new_risk.doenca_pernas = verificar_form_risk(
                request.POST.get('doenca_pernas'))
            new_risk.hipercolestrol = verificar_form_risk(
                request.POST.get('hipercolestrol'))
            new_risk.doenca_cognitiva = verificar_form_risk(
                request.POST.get('doenca_cognitiva'))
            new_risk.comentario = verificar_form_risk(
                request.POST.get('comentario'))
            new_risk.recomendacoes = verificar_form_risk(
                request.POST.get('recomendacoes'))
            new_risk.hemoglobina_gliciada = 0.0
            # new_risk.hemoglobina_gliciada = new_risk.hemoglobina_gliciada.replace(',', '.')
            if 150 > (new_risk.colestrol_nao_hdl) and (new_risk.colestrol_nao_hdl) >= 0:
                new_risk.hemoglobina_gliciada = 3.0
            elif 200 > (new_risk.colestrol_nao_hdl) and (new_risk.colestrol_nao_hdl) >= 150:
                new_risk.hemoglobina_gliciada = 4.0
            elif 250 > (new_risk.colestrol_nao_hdl) and (new_risk.colestrol_nao_hdl) >= 200:
                new_risk.hemoglobina_gliciada = 5.0
            elif new_risk.colestrol_nao_hdl >= 250:
                new_risk.hemoglobina_gliciada = 6.0

            if (request.POST.get('sexo') == 'Feminino' and boolean_idade and boolean_pressao):
                risco = risk_json(file_path_women, new_risk.fumador, new_risk.idade, float(
                    new_risk.hemoglobina_gliciada), new_risk.pressao_arterial)
                new_risk.risco_de_enfarte = risco
                guardar_genero = 'F'

            elif (request.POST.get('sexo') == 'Masculino' and boolean_idade and boolean_pressao):
                risco = risk_json(file_path_men, new_risk.fumador, new_risk.idade, float(
                    new_risk.hemoglobina_gliciada), new_risk.pressao_arterial)
                new_risk.risco_de_enfarte = risco
                guardar_genero = 'M'
            else:
                new_risk.risco_de_enfarte = 0

            new_risk.parteDoUtilizador = parteDoUtilizador
            new_risk.concluido = True
            r.part.concluido = True
            r.save()
            new_risk.imc = calcular_imc(new_risk.peso, new_risk.altura)

            if existing_risk is None:
                # cria uma nova resposta
                r.increment_statistics(f'{part_id}', f'{area_id}', f'{instrument_id}', f'{dimension_id}',
                                       f'{section_id}')
                new_risk.save()
                ris = new_risk
            else:
                # modifica a resposta existente
                existing_risk.peso = float(request.POST.get('peso'))
                # outro
                existing_risk.idade = new_risk.idade
                existing_risk.sexo = new_risk.sexo
                existing_risk.altura = new_risk.altura
                existing_risk.pressao_arterial = new_risk.pressao_arterial
                existing_risk.colestrol_total = new_risk.colestrol_total
                existing_risk.colestrol_hdl = new_risk.colestrol_hdl
                existing_risk.colestrol_nao_hdl = int(
                    new_risk.colestrol_nao_hdl)
                existing_risk.tg = new_risk.tg
                existing_risk.ldl = new_risk.ldl
                existing_risk.cholhdl = float(new_risk.cholhdl)
                existing_risk.eag_hba1 = float(new_risk.eag_hba1)
                existing_risk.ngsp_hba1 = float(new_risk.ngsp_hba1)
                existing_risk.ifcc_hba1 = float(new_risk.ifcc_hba1)
                existing_risk.fumador = new_risk.fumador
                existing_risk.diabetes = new_risk.diabetes
                existing_risk.pre_diabetico = new_risk.pre_diabetico
                existing_risk.pergunta_cardiovascular = new_risk.pergunta_cardiovascular
                existing_risk.pat_id = new_risk.pat_id
                existing_risk.batimentos = new_risk.batimentos
                existing_risk.hemoglobina_gliciada = 0.0
                # new_risk.hemoglobina_gliciada = new_risk.hemoglobina_gliciada.replace(',', '.')
                if 150 > (existing_risk.colestrol_nao_hdl) and (existing_risk.colestrol_nao_hdl) >= 0:
                    existing_risk.hemoglobina_gliciada = 3.0
                elif 200 > (existing_risk.colestrol_nao_hdl) and (existing_risk.colestrol_nao_hdl) >= 150:
                    existing_risk.hemoglobina_gliciada = 4.0
                elif 250 > (existing_risk.colestrol_nao_hdl) and (existing_risk.colestrol_nao_hdl) >= 200:
                    existing_risk.hemoglobina_gliciada = 5.0
                elif (existing_risk.colestrol_nao_hdl) >= 250:
                    existing_risk.hemoglobina_gliciada = 6.0
                existing_risk.anos_diabetes = verificar_form_risk(
                    request.POST.get('anos_diabetes'))
                existing_risk.pre_diabetico = verificar_form_risk(
                    request.POST.get('pre_diabetico'))
                existing_risk.pergunta_cardiovascular = verificar_form_risk(
                    request.POST.get('pergunta_cardiovascular'))
                existing_risk.horas_jejum = verificar_form_risk(
                    request.POST.get('horas_jejum'))
                existing_risk.doenca_cognitiva = verificar_form_risk(
                    request.POST.get('doenca_cognitiva'))
                existing_risk.avc = verificar_form_risk(
                    request.POST.get('avc'))
                existing_risk.enfarte = verificar_form_risk(
                    request.POST.get('enfarte'))
                existing_risk.doenca_rins = verificar_form_risk(
                    request.POST.get('doenca_rins'))
                existing_risk.doenca_pernas = verificar_form_risk(
                    request.POST.get('doenca_pernas'))
                existing_risk.hipercolestrol = verificar_form_risk(
                    request.POST.get('hipercolestrol'))
                existing_risk.comentario = verificar_form_risk(
                    request.POST.get('comentario'))
                existing_risk.recomendacoes = verificar_form_risk(
                    request.POST.get('recomendacoes'))
                existing_risk.imc = calcular_imc(
                    existing_risk.peso, existing_risk.altura)
                if (request.POST.get('sexo') == 'Feminino' and boolean_idade and boolean_pressao):
                    risco = risk_json(file_path_women, existing_risk.fumador, existing_risk.idade, float(
                        existing_risk.hemoglobina_gliciada), existing_risk.pressao_arterial)
                    existing_risk.risco_de_enfarte = risco
                    guardar_genero = 'F'
                elif request.POST.get('sexo') == 'Masculino' and boolean_idade and boolean_pressao:
                    risco = risk_json(file_path_men, existing_risk.fumador, existing_risk.idade, float(
                        existing_risk.hemoglobina_gliciada), existing_risk.pressao_arterial)
                    existing_risk.risco_de_enfarte = risco
                    guardar_genero = 'M'
                else:
                    existing_risk.risco_de_enfarte = 0
                existing_risk.save()
                ris = existing_risk
            r.change_quotation(f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}',
                               new_risk.risco_de_enfarte)
            username = request.user.username

            concluido_risk = True

            gera_relatorio_risk_pdf(
                ris, patient, username, guardar_genero, boolean_idade, boolean_pressao)
            guardar_genero = ""
            # este context serve para desaparecer o breadcrumb do risk
            context['concluido_risk'] = concluido_risk

            if is_cuidador == 'is_cuidador':
                return redirect('protocolo:profile_cuidador', cuidador_id=patient_id)
            else:
                return redirect('protocolo:participant', participant_id=patient_id)

        if question.question_type == 3:
            return redirect('protocolo:instruments', protocol_id=protocol_id, part_id=part_id, area_id=area_id,
                            patient_id=patient_id, is_cuidador=is_cuidador)
        elif question.name == "Relação com o Avaliador" or question.name == "Cooperação dada na entrevista" or question.name == "Questionário Sociodemográfico":
            return redirect('protocolo:areas', protocol_id=protocol_id, part_id=part_id, patient_id=patient_id, is_cuidador=is_cuidador)
        elif question.section.dimension.name == "None" and question.section.name == "None":
            return redirect('protocolo:instruments', protocol_id=protocol_id, part_id=part_id, area_id=area_id,
                            patient_id=patient_id, is_cuidador=is_cuidador)
        elif question.section.name == "None" or question.question_type == 9 or question.section.dimension.number_of_questions == 1:
            return redirect('protocolo:dimensions', protocol_id=protocol_id, part_id=part_id, area_id=area_id,
                            instrument_id=instrument_id, patient_id=patient_id, is_cuidador=is_cuidador)
        else:
            return redirect('protocolo:sections', protocol_id=protocol_id, part_id=part_id, area_id=area_id,
                            instrument_id=instrument_id, dimension_id=dimension_id, patient_id=patient_id, is_cuidador=is_cuidador)

    end = time.time()
    print("Question", (end - start))

    return render(request, 'protocolo/question.html', context)


@login_required(login_url='login')
def delete_button(request, partedoutilizador_id, participant_id):

    item = ParteDoUtilizador.objects.get(pk=partedoutilizador_id)
    print("item:")
    print(item)
    print("fora do item")
    item.delete()
    return redirect('protocolo:participant', participant_id=participant_id)


@login_required(login_url='login')
def report_view(request, resolution_id):
    start = time.time()
    r = Resolution.objects.get(pk=resolution_id)
    nome_parte = r.part.part.name
    areas = Area.objects.filter(part=r.part.part)
    report_json = r.statistics
    report_json_dumps = json.dumps(
        report_json, indent=1, sort_keys=False, ensure_ascii=False)
    reporte = Report.objects.filter(resolution=r).get()
    report = {}
    answers = Answer.objects.filter(
        resolution=r).order_by("question__section__order")
    done = []
    chc = []
    nr_areas = len(areas) - 3
    nr_total_instrumentos = -3

    report_obj = Report.objects.get(resolution=r)
    report_obj.refresh_report(answers)
    # inicializar instruments a nada
    instruments = None

    for area in areas.order_by('order'):
        nr_total_instrumentos += area.number_of_instruments
        if area.name != 'Cognição':
            report[area.name] = {
                'nr_instrumentos': area.number_of_instruments,
            }
        instruments = Instrument.objects.all().order_by('order').filter(area=area)

        for instrument in instruments:
            if instrument.name != None:
                if 'ABVD' in instrument.name:
                    respondido = r.statistics.get(f'{area.id}').get(
                        f'{instrument.id}').get('answered') > 0
                    if respondido:
                        report[area.name]['nr_instrumentos'] += 1
                    names = ['Atividades Corporais', 'Atividades Motoras',
                             'Atividades Mentais', 'Atividades Sensoriais']
                    quotations = [report_obj.abvd_atividades_corporais_quotation, report_obj.abvd_atividades_motoras_quotation,
                                  report_obj.abvd_atividades_mentais_quotation, report_obj.abvd_atividades_sensoriais_quotation]
                    total = sum(quotations)

                    dimensao = Dimension.objects.filter(instrument=instrument)
                    dimensao_acti_corp = dimensao.filter(
                        name='Atividades Corporais').get()
                    dimensao_acti_moto = dimensao.filter(
                        name='Atividades Motoras').get()
                    dimensao_acti_ment = dimensao.filter(
                        name='Atividades Mentais').get()
                    dimensao_acti_sens = dimensao.filter(
                        name='Atividades Sensoriais').get()

                    done.append(instrument.name)

                    print(
                        f"Creating Graph for {instrument.name} ; h-max={instrument.highest_max_quotation}, min={instrument.minimum_quotation}")
                    svg, png = make_graph(
                        names, quotations, 0, instrument.highest_max_quotation)
                    numero_abvd = instrument.highest_max_quotation
                    report[area.name][instrument.name] = {
                        'max_ati_corp': dimensao_acti_corp.maximum_quotation,
                        'max_ati_moto': dimensao_acti_moto.maximum_quotation,
                        'max_ati_ment': dimensao_acti_ment.maximum_quotation,
                        'max_ati_sens': dimensao_acti_sens.maximum_quotation,
                        'evaluation': report_obj.abvd_evaluation,
                        'atividades_corporais': report_obj.abvd_atividades_corporais_quotation,
                        'atividades_motoras': report_obj.abvd_atividades_motoras_quotation,
                        'atividades_mentais': report_obj.abvd_atividades_mentais_quotation,
                        'atividades_sensoriais': report_obj.abvd_atividades_sensoriais_quotation,
                        'total': total,
                        'graph': svg,
                        'img': png,
                        'respondido': respondido,
                        'notes': 'abvd_notes',
                    }
                    print(f'{dimensao_acti_corp.maximum_quotation}')

                elif 'AIVD' in instrument.name:
                    respondido = r.statistics.get(f'{area.id}').get(
                        f'{instrument.id}').get('answered') > 0
                    if respondido:
                        report[area.name]['nr_instrumentos'] += 1
                    report[area.name][instrument.name] = {
                        'evaluation': report_obj.aivd_evaluation,
                        'utilizacao_telefone': report_obj.aivd_utilizacao_telefone_quotation,
                        'fazer_compras': report_obj.aivd_fazer_compras_quotation,
                        'preparar_refeicoes': report_obj.aivd_preparacao_refeicoes_quotation,
                        'tarefas_domesticas': report_obj.aivd_tarefas_domesticas_quotation,
                        'lavar_roupa': report_obj.aivd_lavar_roupa_quotation,
                        'utilizar_transportes': report_obj.aivd_utilizar_transportes_quotation,
                        'manejo_medicacao': report_obj.aivd_manejo_mediacao_quotation,
                        'responsabilidade_financeira': report_obj.aivd_responsabilidades_financeiras_quotation,
                        'respondido': False,
                        'respondido': respondido,
                        'notes': 'aivd_notes',
                    }

                elif 'BSI' in instrument.name:
                    respondido = r.statistics.get(f'{area.id}').get(
                        f'{instrument.id}').get('answered') > 0
                    if respondido:
                        report[area.name]['nr_instrumentos'] += 1
                    names = ['Somatização', 'Obsessão-Compulsão', 'Sensibilidade Interpessoal', 'Depressão',
                             'Ansiedade', 'Hostilidade', 'Ansiedade Fóbica', 'Ideação Paranóide', 'Psicoticismo']
                    quotations = [report_obj.bsi_somatizacao, report_obj.bsi_obssessivo_compulsivo, report_obj.bsi_sensibilidade_interpessoal, report_obj.bsi_depressao,
                                  report_obj.bsi_ansiedade, report_obj.bsi_hostilidade, report_obj.bsi_ansiedade_fobica, report_obj.bsi_paranoide, report_obj.bsi_psicoticismo]
                    done.append(instrument.name)
                    print(
                        f"Creating Graph for {instrument.name} ; h-max={instrument.highest_max_quotation}, min={instrument.minimum_quotation}")
                    dimensao = Dimension.objects.filter(instrument=instrument)
                    svg, png = make_graph(names, quotations, 0, 28)

                    # Valores do max têm de ser hardcoded
                    report[area.name][instrument.name] = {

                        'somatizacao': report_obj.bsi_somatizacao,
                        'max_somatizacao': 0,
                        'obsessao': report_obj.bsi_obssessivo_compulsivo,
                        'max_obsessao': 0,
                        'depressao': report_obj.bsi_depressao,
                        'max_depressao': 0,
                        'sensibilidade_interpessoal': report_obj.bsi_sensibilidade_interpessoal,
                        'max_sensibilidade': 0,
                        'ansiedade': report_obj.bsi_ansiedade,
                        'max_ansiedade': 0,
                        'hostilidade': report_obj.bsi_hostilidade,
                        'max_hostilidade': 0,
                        'ansiedade_fobica': report_obj.bsi_ansiedade_fobica,
                        'max_ansiedade': 0,
                        'paranoide': report_obj.bsi_paranoide,
                        'max_paranoide': 0,
                        'psicoticismo': report_obj.bsi_psicoticismo,
                        'max_psicoticismo': 0,
                        'igs': round(report_obj.bsi_igs, 2),
                        'tsp': round(report_obj.bsi_tsp, 2),
                        'isp': round(report_obj.bsi_isp, 2),
                        'graph': svg,
                        'img': png,
                        'respondido': respondido,
                        'notes': 'bsi_notes',
                    }

                elif 'ACE-R' in instrument.name:
                    respondido = r.statistics.get(f'{area.id}').get(
                        f'{instrument.id}').get('answered') > 0
                    if respondido:
                        report[area.name]['nr_instrumentos'] += 1
                    names = ['Atenção e Orientação', 'Memória',
                             'Fluência', 'Linguagem', 'Visuo-Espacial']
                    quotations = [report_obj.acer_atencao_orientacao_quotation, report_obj.acer_memoria_quotation,
                                  report_obj.acer_fluencia_quotation, report_obj.acer_linguagem_quotation, report_obj.acer_visuo_espacial_quotation]
                    done.append(instrument.name)
                    print(
                        f"Creating Graph for {instrument.name} ; h-max={instrument.highest_max_quotation}, min={instrument.minimum_quotation}")
                    dimensao = Dimension.objects.filter(instrument=instrument)
                    dimensao_aten_orient = dimensao.filter(
                        name='Atenção e Orientação').get()
                    dimensao_memoria = dimensao.filter(name='Memória').get()
                    dimensao_fluencia = dimensao.filter(name='Fluência').get()
                    dimensao_linguagem = dimensao.filter(
                        name='Linguagem').get()
                    dimensao_visuo_espacial = dimensao.filter(
                        name='Visuo-Espacial').get()
                    print(dimensao_aten_orient.maximum_quotation)
                    svg, png = make_graph(
                        names, quotations, 0, instrument.highest_max_quotation)
                    report[area.name][instrument.name] = {
                        'total': sum(quotations),
                        'max_aten_orient': dimensao_aten_orient.maximum_quotation,
                        'max_memoria': dimensao_memoria.maximum_quotation,
                        'max_fluencia': dimensao_fluencia.maximum_quotation,
                        'max_linguagem': dimensao_linguagem.maximum_quotation,
                        'max_visuo_espacial': dimensao_visuo_espacial.maximum_quotation,
                        'evaluation': report_obj.acer_evaluation,
                        'acer_atencao_orientacao': report_obj.acer_atencao_orientacao_quotation,
                        'acer_memoria': report_obj.acer_memoria_quotation,
                        'acer_fluencia': report_obj.acer_fluencia_quotation,
                        'acer_linguagem': report_obj.acer_linguagem_quotation,
                        'acer_visuo_espacial': report_obj.acer_visuo_espacial_quotation,
                        'graph': svg,
                        'img': png,
                        'respondido': respondido,
                        'notes': 'acer_notes',
                    }

                elif 'MMSE' in instrument.name:
                    respondido = r.statistics.get(f'{area.id}').get(
                        f'{instrument.id}').get('answered') > 0
                    if respondido:
                        report[area.name]['nr_instrumentos'] += 1
                    names = ['Atenção e Orientação', 'Memória',
                             'Linguagem', 'Visuo-Espacial']
                    quotations = [report_obj.mmse_atencao_orientacao_quotation, report_obj.mmse_memoria_quotation,
                                  report_obj.mmse_lingua_quotation, report_obj.mmse_visuo_espacial_quotation]
                    total = sum(quotations)
                    done.append(instrument.name)
                    print(
                        f"Creating Graph for {instrument.name} ; h-max={instrument.highest_max_quotation}, min={instrument.minimum_quotation}")
                    dimensao = Dimension.objects.filter(instrument=instrument)
                    dimensao_atent_orient = dimensao.filter(
                        name='Atenção e Orientação').get()
                    dimensao_memoria = dimensao.filter(name='Memória').get()
                    dimensao_linguagem = dimensao.filter(
                        name='Linguagem').get()
                    dimensao_visuo_espacial = dimensao.filter(
                        name='Visuo-Espacial').get()
                    svg, png = make_graph(
                        names, quotations, 0, instrument.highest_max_quotation)
                    report[area.name][instrument.name] = {
                        'total': total,
                        'evaluation': report_obj.mmse_evaluation,
                        'max_atent_orient': dimensao_atent_orient.maximum_quotation,
                        'max_memoria': dimensao_memoria.maximum_quotation,
                        'max_linguagem': dimensao_linguagem.maximum_quotation,
                        'max_visuo_espacial': dimensao_visuo_espacial.maximum_quotation,
                        'mmse_atencao_orientacao': report_obj.mmse_atencao_orientacao_quotation,
                        'mmse_memoria': report_obj.mmse_memoria_quotation,
                        'mmse_linguagem': report_obj.mmse_lingua_quotation,
                        'mmse_visuo_espacial': report_obj.mmse_visuo_espacial_quotation,
                        'graph': svg,
                        'img': png,
                        'respondido': respondido,
                        'notes': 'mmse_notes',
                    }

                elif 'PANAS' in instrument.name:
                    respondido = r.statistics.get(f'{area.id}').get(
                        f'{instrument.id}').get('answered') > 0
                    if respondido:
                        report[area.name]['nr_instrumentos'] += 1
                    done.append(instrument.name)
                    report[area.name][instrument.name] = {
                        'interessado': report_obj.panas_interessado,
                        'nervoso': report_obj.panas_nervoso,
                        'entusiasmado': report_obj.panas_entusiasmado,
                        'amedrontado': report_obj.panas_amedrontado,
                        'inspirado': report_obj.panas_inspirado,
                        'ativo': report_obj.panas_ativo,
                        'assustado': report_obj.panas_assustado,
                        'culpado': report_obj.panas_culpado,
                        'determinado': report_obj.panas_determinado,
                        'atormentado': report_obj.panas_atormentado,
                        'respondido': respondido,
                        'notes': 'panas_notes',
                    }

                elif 'HADS' in instrument.name:
                    respondido = r.statistics.get(f'{area.id}').get(
                        f'{instrument.id}').get('answered') > 0
                    if respondido:
                        report[area.name]['nr_instrumentos'] += 1
                    done.append(instrument.name)
                    report[area.name][instrument.name] = {
                        'ansiedade_evaluation': report_obj.hads_estado_ansiedade_evaluation,
                        'ansiedade_quotation': report_obj.hads_estado_ansiedade_quotation,
                        'depressao_evaluation': report_obj.hads_estado_depressao_evaluation,
                        'depressao_quotation': report_obj.hads_estado_depressao_quotation,
                        'respondido': respondido,
                        'notes': 'hads_notes',
                    }

                elif 'GDS' in instrument.name:
                    respondido = r.statistics.get(f'{area.id}').get(
                        f'{instrument.id}').get('answered') > 0
                    if respondido:
                        report[area.name]['nr_instrumentos'] += 1
                    done.append(instrument.name)
                    report[area.name][instrument.name] = {
                        'gds': report_obj.gds_nivel,
                        'respondido': respondido,
                        'notes': 'gds_notes',
                    }

                elif 'Áreas Complementares' in instrument.name:
                    respondido = r.statistics.get(f'{area.id}').get(
                        f'{instrument.id}').get('answered') > 0
                    if respondido:
                        report[area.name]['nr_instrumentos'] += 1
                    done.append(instrument.name)
                    report[area.name][instrument.name] = {
                        'memoria_visual_imediata': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('21').get('quotation'),
                        'memoria_visual_diferida': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('22').get('quotation'),
                        'atencao_mantida': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('23').get('quotation'),
                        'atencao_dividida': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('24').get('quotation'),
                        'orientacao_esq_dir': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('25').get('quotation'),
                        'abstracao_verbal': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('26').get('quotation'),
                        'compreensao_instrucoes': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('27').get('quotation'),
                        'respondido': respondido,
                        'notes': 'ac_notes',
                    }

                elif 'None' in instrument.name:
                    respondido = r.statistics.get(f'{area.id}').get(
                        f'{instrument.id}').get('answered') > 0
                    if respondido:
                        report[area.name]['nr_instrumentos'] += 1
                    if area.name == 'Consciência, Humor e Comportamento':
                        cons = MultipleChoicesCheckbox.objects.filter(
                            answer__resolution=r, answer__question__name='Consciência')
                        cons_list = [self.choice.name for self in cons]
                        chc.append(cons_list)
                        mot = MultipleChoicesCheckbox.objects.filter(
                            answer__resolution=r, answer__question__name='Atividade Motora')
                        mot_list = [self.choice.name for self in mot]
                        chc.append(mot_list)
                        hum = MultipleChoicesCheckbox.objects.filter(
                            answer__resolution=r, answer__question__name='Humor')
                        hum_list = [self.choice.name for self in hum]
                        chc.append(hum_list)
                        report[area.name][area.name] = {
                            'chc_consciencia': ", ".join(cons_list),
                            'chc_motora': ", ".join(mot_list),
                            'chc_humor': ", ".join(hum_list),
                            'respondido': respondido,
                            'notes': 'chc_notes',
                        }

                    elif area.name == 'Cooperação dada na entrevista':
                        coop = ''
                        respondido = r.statistics.get(f'{area.id}').get(
                            f'{instrument.id}').get('answered') > 0
                        report[area.name][area.name] = {
                            'respondido': respondido,
                            'notes': 'coop_notes',
                        }
                        if respondido:
                            coop = Answer.objects.filter(
                                resolution=r, question__name='Cooperação dada na entrevista').get().multiple_choice_answer.name
                            report[area.name][area.name]['coop'] = coop

                    elif area.name == 'Relação com o Avaliador':
                        rel = ''
                        respondido = r.statistics.get(f'{area.id}').get(
                            f'{instrument.id}').get('answered') > 0
                        report[area.name][area.name] = {
                            'respondido': respondido,
                            'notes': 'rel_notes',
                        }
                        if respondido:
                            rel = Answer.objects.filter(
                                resolution=r, question__name='Relação com o Avaliador').get().multiple_choice_answer.name
                            report[area.name][area.name]['rel'] = rel

    if 'Cognição' in report[area.name]:
        print('dividir')
        report['Cognição']['nr_instrumentos'] /= 2

    context = {
        'report_json': report_json,
        'report_json_dumps': report_json_dumps,
        'report': report,
        'reporte': reporte,
        'resolution': r,
        'answers': answers,
        'instruments': Instrument.objects.all(),
        'questions': Question.objects.all(),
        'areas': areas,
        'nome_parte': nome_parte,
        'nr_areas': nr_areas,
        'nr_total_instrumentos': nr_total_instrumentos*2,
    }

    end = time.time()
    print("Report", (end - start))

    return render(request, 'protocolo/report2.html', context)


@login_required(login_url='login')
def report2(request, resolution_id):
    start = time.time()
    r = Resolution.objects.get(pk=resolution_id)
    nome_parte = r.part.part.name
    areas = Area.objects.filter(part=r.part.part)
    report_json = r.statistics
    report_json_dumps = json.dumps(
        report_json, indent=1, sort_keys=False, ensure_ascii=False)
    reporte = Report.objects.filter(resolution=r).get()
    report = {}
    answers = Answer.objects.filter(
        resolution=r).order_by("question__section__order")
    done = []
    chc = []
    nr_areas = len(areas) - 3
    nr_total_instrumentos = -3

    report_obj = Report.objects.get(resolution=r)
    report_obj.refresh_report(answers)
    # inicializar instruments a nada
    instruments = None

    for area in areas.order_by('order'):
        nr_total_instrumentos += area.number_of_instruments
        report[area.name] = {
            'nr_instrumentos': area.number_of_instruments,
        }
        instruments = Instrument.objects.all().order_by('order').filter(area=area)

        for instrument in instruments:
            if instrument.name != None:
                if 'ABVD' in instrument.name:
                    # abvd = Instrument.objects.get(name='ABVD')
                    names = ['Atividades Corporais', 'Atividades Motoras',
                             'Atividades Mentais', 'Atividades Sensoriais']
                    quotations = [report_obj.abvd_atividades_corporais_quotation, report_obj.abvd_atividades_motoras_quotation,
                                  report_obj.abvd_atividades_mentais_quotation, report_obj.abvd_atividades_sensoriais_quotation]
                    total = sum(quotations)
                    done.append(instrument.name)
                    print(
                        f"Creating Graph for {instrument.name} ; h-max={instrument.highest_max_quotation}, min={instrument.minimum_quotation}")
                    svg, png = make_graph(
                        names, quotations, 0, instrument.highest_max_quotation)
                    report[area.name][instrument.name] = {
                        'evaluation': report_obj.abvd_evaluation,
                        'atividades_corporais': report_obj.abvd_atividades_corporais_quotation,
                        'atividades_motoras': report_obj.abvd_atividades_motoras_quotation,
                        'atividades_mentais': report_obj.abvd_atividades_mentais_quotation,
                        'atividades_sensoriais': report_obj.abvd_atividades_sensoriais_quotation,
                        'total': total,
                        'graph': svg,
                        'img': png,
                        'respondido': respondido,
                        'notes': 'abvd_notes',
                    }

                elif 'AIVD' in instrument.name:
                    report[area.name][instrument.name] = {
                        'evaluation': report_obj.aivd_evaluation,
                        'utilizacao_telefone': report_obj.aivd_utilizacao_telefone_quotation,
                        'fazer_compras': report_obj.aivd_fazer_compras_quotation,
                        'preparar_refeicoes': report_obj.aivd_preparacao_refeicoes_quotation,
                        'tarefas_domesticas': report_obj.aivd_tarefas_domesticas_quotation,
                        'lavar_roupa': report_obj.aivd_lavar_roupa_quotation,
                        'utilizar_transportes': report_obj.aivd_utilizar_transportes_quotation,
                        'manejo_medicacao': report_obj.aivd_manejo_mediacao_quotation,
                        'responsabilidade_financeira': report_obj.aivd_responsabilidades_financeiras_quotation,
                        'respondido': False,
                        'respondido': respondido,
                        'notes': 'aivd_notes',
                    }

                elif 'BSI' in instrument.name:
                    names = ['Somatização', 'Obsessão-Compulsão', 'Sensibilidade Interpessoal', 'Depressão',
                             'Ansiedade', 'Hostilidade', 'Ansiedade Fóbica', 'Ideação Paranóide', 'Psicoticismo']
                    quotations = [report_obj.bsi_somatizacao, report_obj.bsi_obssessivo_compulsivo, report_obj.bsi_sensibilidade_interpessoal, report_obj.bsi_depressao,
                                  report_obj.bsi_ansiedade, report_obj.bsi_hostilidade, report_obj.bsi_ansiedade_fobica, report_obj.bsi_paranoide, report_obj.bsi_psicoticismo]
                    done.append(instrument.name)
                    print(
                        f"Creating Graph for {instrument.name} ; h-max={instrument.highest_max_quotation}, min={instrument.minimum_quotation}")
                    svg, png = make_graph(names, quotations, 0, 28)
                    numero_max_bsi = instrument.highest_max_quotation
                    report[area.name][instrument.name] = {
                        'somatizacao': report_obj.bsi_somatizacao,
                        'obsessao': report_obj.bsi_obssessivo_compulsivo,
                        'depressao': report_obj.bsi_depressao,
                        'sensibilidade_interpessoal': report_obj.bsi_sensibilidade_interpessoal,
                        'ansiedade': report_obj.bsi_ansiedade,
                        'hostilidade': report_obj.bsi_hostilidade,
                        'ansiedade_fobica': report_obj.bsi_ansiedade_fobica,
                        'paranoide': report_obj.bsi_paranoide,
                        'psicoticismo': report_obj.bsi_psicoticismo,
                        'igs': round(report_obj.bsi_igs, 2),
                        'tsp': round(report_obj.bsi_tsp, 2),
                        'isp': round(report_obj.bsi_isp, 2),
                        'graph': svg,
                        'img': png,
                        'respondido': respondido,
                        'notes': 'bsi_notes',
                    }

                elif 'ACE-R' in instrument.name:
                    names = ['Atenção e Orientação', 'Memória',
                             'Fluência', 'Linguagem', 'Visuo-Espacial']
                    quotations = [report_obj.acer_atencao_orientacao_quotation, report_obj.acer_memoria_quotation,
                                  report_obj.acer_fluencia_quotation, report_obj.acer_linguagem_quotation, report_obj.acer_visuo_espacial_quotation]
                    done.append(instrument.name)
                    print(
                        f"Creating Graph for {instrument.name} ; h-max={instrument.highest_max_quotation}, min={instrument.minimum_quotation}")
                    svg, png = make_graph(
                        names, quotations, 0, instrument.highest_max_quotation)
                    numero_max_acer = instrument.highest_max_quotation
                    report[area.name][instrument.name] = {
                        'total': sum(quotations),
                        'evaluation': report_obj.acer_evaluation,
                        'acer_atencao_orientacao': report_obj.acer_atencao_orientacao_quotation,
                        'acer_memoria': report_obj.acer_memoria_quotation,
                        'acer_fluencia': report_obj.acer_fluencia_quotation,
                        'acer_linguagem': report_obj.acer_linguagem_quotation,
                        'acer_visuo_espacial': report_obj.acer_visuo_espacial_quotation,
                        'graph': svg,
                        'img': png,
                        'respondido': respondido,
                        'notes': 'acer_notes',
                    }
                elif 'MMSE' in instrument.name:
                    names = ['Atenção e Orientação', 'Memória',
                             'Linguagem', 'Visuo-Espacial']
                    quotations = [report_obj.mmse_atencao_orientacao_quotation, report_obj.mmse_memoria_quotation,
                                  report_obj.mmse_lingua_quotation, report_obj.mmse_visuo_espacial_quotation]
                    total = sum(quotations)
                    done.append(instrument.name)
                    print(
                        f"Creating Graph for {instrument.name} ; h-max={instrument.highest_max_quotation}, min={instrument.minimum_quotation}")
                    svg, png = make_graph(
                        names, quotations, 0, instrument.highest_max_quotation)
                    numero_max_mmse = instrument.highest_max_quotation
                    report[area.name][instrument.name] = {
                        'total': total,
                        'evaluation': report_obj.mmse_evaluation,
                        'mmse_atencao_orientacao': report_obj.mmse_atencao_orientacao_quotation,
                        'mmse_memoria': report_obj.mmse_memoria_quotation,
                        'mmse_linguagem': report_obj.mmse_lingua_quotation,
                        'mmse_visuo_espacial': report_obj.mmse_visuo_espacial_quotation,
                        'graph': svg,
                        'img': png,
                        'respondido': respondido,
                        'notes': 'mmse_notes',
                    }

                elif 'PANAS' in instrument.name:
                    done.append(instrument.name)
                    report[area.name][instrument.name] = {
                        'interessado': report_obj.panas_interessado,
                        'nervoso': report_obj.panas_nervoso,
                        'entusiasmado': report_obj.panas_entusiasmado,
                        'amedrontado': report_obj.panas_amedrontado,
                        'inspirado': report_obj.panas_inspirado,
                        'ativo': report_obj.panas_ativo,
                        'assustado': report_obj.panas_assustado,
                        'culpado': report_obj.panas_culpado,
                        'determinado': report_obj.panas_determinado,
                        'atormentado': report_obj.panas_atormentado,
                        'respondido': respondido,
                        'notes': 'panas_notes',
                    }

                elif 'HADS' in instrument.name:
                    done.append(instrument.name)
                    report[area.name][instrument.name] = {
                        'ansiedade_evaluation': report_obj.hads_estado_ansiedade_evaluation,
                        'ansiedade_quotation': report_obj.hads_estado_ansiedade_quotation,
                        'depressao_evaluation': report_obj.hads_estado_depressao_evaluation,
                        'depressao_quotation': report_obj.hads_estado_depressao_quotation,
                        'respondido': respondido,
                        'notes': 'hads_notes',
                    }

                elif 'GDS' in instrument.name:
                    done.append(instrument.name)
                    report[area.name][instrument.name] = {
                        'gds': report_obj.gds_nivel,
                        'respondido': respondido,
                        'notes': 'gds_notes',
                    }

                elif 'Áreas Complementares' in instrument.name:
                    done.append(instrument.name)
                    report[area.name][instrument.name] = {
                        'memoria_visual_imediata': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('21').get('quotation'),
                        'memoria_visual_diferida': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('22').get('quotation'),
                        'atencao_mantida': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('23').get('quotation'),
                        'atencao_dividida': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('24').get('quotation'),
                        'orientacao_esq_dir': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('25').get('quotation'),
                        'abstracao_verbal': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('26').get('quotation'),
                        'compreensao_instrucoes': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('27').get('quotation'),
                        'respondido': respondido,
                        'notes': 'ac_notes',
                    }

                elif 'None' in instrument.name:
                    if area.name == 'Consciência, Humor e Comportamento':
                        cons = MultipleChoicesCheckbox.objects.filter(
                            answer__resolution=r, answer__question__name='Consciência')
                        cons_list = [self.choice.name for self in cons]
                        chc.append(cons_list)
                        mot = MultipleChoicesCheckbox.objects.filter(
                            answer__resolution=r, answer__question__name='Atividade Motora')
                        mot_list = [self.choice.name for self in mot]
                        chc.append(mot_list)
                        hum = MultipleChoicesCheckbox.objects.filter(
                            answer__resolution=r, answer__question__name='Humor')
                        hum_list = [self.choice.name for self in hum]
                        chc.append(hum_list)
                        report[area.name][area.name] = {
                            'chc_consciencia': ", ".join(cons_list),
                            'chc_motora': ", ".join(mot_list),
                            'chc_humor': ", ".join(hum_list),
                            'respondido': respondido,
                            'notes': 'chc_notes',
                        }

                    elif area.name == 'Cooperação dada na entrevista':
                        coop = ''
                        respondido = r.statistics.get(f'{area.id}').get(
                            f'{instrument.id}').get('answered') > 0
                        report[area.name][area.name] = {
                            'respondido': respondido,
                            'notes': 'coop_notes',
                        }
                        if respondido:
                            coop = Answer.objects.filter(
                                resolution=r, question__name='Cooperação dada na entrevista').get().multiple_choice_answer.name
                            report[area.name][area.name]['coop'] = coop

                    elif area.name == 'Relação com o Avaliador':
                        rel = ''
                        respondido = r.statistics.get(f'{area.id}').get(
                            f'{instrument.id}').get('answered') > 0
                        report[area.name][area.name] = {
                            'respondido': respondido,
                            'notes': 'rel_notes',
                        }
                        if respondido:
                            rel = Answer.objects.filter(
                                resolution=r, question__name='Relação com o Avaliador').get().multiple_choice_answer.name
                            report[area.name][area.name]['rel'] = rel

    gera_relatorio_parte(r, chc, coop, rel, reporte, report)

    context = {'report_json': report_json,
               'report_json_dumps': report_json_dumps,
               'report': report,
               'reporte': reporte,
               'resolution': r,
               'answers': answers,
               'instruments': Instrument.objects.all(),
               'questions': Question.objects.all(),
               'areas': areas,
               'nome_parte': nome_parte,
               'nr_areas': nr_areas,
               'nr_total_instrumentos': nr_total_instrumentos*2,
               }

    return render(request, 'protocolo/report2.html', context)


@login_required(login_url='login')
def report_risk(request):
    start = time.time()
    lines = []

    risk = Risk.objects.all()
    for r in risk:
        lines.append("Relatório de Risco: ")
        lines.append("Idade:" + str(r.idade))
        lines.append("Sexo:" + str(r.sexo))
        lines.append("Colesterol:" + str(r.colestrol_total))
        lines.append("pressao_arterial:" + str(r.pressao_arterial))
        lines.append("O risco atual é de:" + str(r.risco_de_enfarte))
        lines.append("Comentario:" + str(r.comentario))
        lines.append("Recomendações adicionais:" + str(r.recomendacoes))

    context = {'lines': lines
               }
    parts_risk = render(request, 'protocolo/parts_risk.html', context)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="Risco.pdf"'
    pdf = pisa.pisaDocument(parts_risk.content, response)
    if not pdf.err:
        return response

    end = time.time()

    print("Report-risk", (end - start))
    return HttpResponse("Erro ao gerar PDF")


@login_required(login_url='login')
def protocol_participants_view(request, protocol_id):
    doctor = request.user
    protocolo = Protocol.objects.get(pk=protocol_id)
    participants = Participante.objects.filter(avaliador=doctor)
    resolutions = Resolution.objects.filter(doctor=doctor)

    context = {'participants': participants,
               'resolutions': resolutions, 'protocolo': protocolo}
    return render(request, 'protocolo/protocol-participants.html', context)


def get_nr_participantes(user):
    return Participante.objects.filter(avaliador=user).count()


@login_required(login_url='login')
def participants_view(request):
    doctor = request.user
    participants = []
    avaliadores = []
    lista_nr_participantes = []
    try:
        user = Avaliador.objects.get(user=doctor)
        participants = Participante.objects.filter(avaliador=doctor)
        print(participants)

    except Avaliador.DoesNotExist:
        try:
            user = Administrador.objects.get(user=doctor)
            participants = Participante.objects.filter(
                referenciacao=user.reference)
            cuidadores = Cuidador.objects.filter(referenciacao=user.reference)
            all_avaliadores = DjangoGroup.objects.get(
                name="Avaliador").user_set.all()

            for avaliador in all_avaliadores:
                if Avaliador.objects.filter(user=avaliador, reference=user.reference).exists():
                    avaliadores.append(avaliador)
                    lista_nr_participantes.append(get_nr_participantes(doctor))
        except Administrador.DoesNotExist:
            user = None

    resolutions = Resolution.objects.filter(doctor=doctor)
    # cuidadores = Cuidador.objects.filter(referenciacao=user.reference)
    cuidadores = Cuidador.objects.filter(avaliador=doctor)

    if request.user.is_superuser:
        participants = Participante.objects.all()
        resolutions = Resolution.objects.all()

    context = {'participants': participants,
               'resolutions': resolutions,
               'avaliadores': avaliadores,
               'avaliadores_zip': zip(avaliadores, lista_nr_participantes),
               'cuidadores': cuidadores, }
    return render(request, 'protocolo/participants.html', context)


def login_view(request):
    next = request.GET.get('next')
    if request.method == 'POST':
        next = request.POST.get('next')
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)

        if user is not None:
            login(request, user)
            # print(request.POST)
            next_url = request.POST.get('next')
            if next_url:
                return HttpResponseRedirect(next_url)
            elif request.user.groups.filter(name='Avaliador').exists():
                return redirect('protocolo:dashboard')
            elif request.user.groups.filter(name__in=['Dinamizador', 'Cuidador', 'Participante']).exists():
                return redirect('protocolo:diario:dashboard_Care')

    context = {
        'next': next,
    }

    return render(request, 'protocolo/login.html')


def logout_view(request):
    logout(request)
    # Mudei isto de protocolo/login.html
    return render(request, 'mentha/base.html')


@login_required(login_url='login')
def profile_view(request, participant_id):
    # Falta mostrar as resoluções das partes feitas no perfil e os seus relatorios

    patient = Participante.objects.filter(pk=participant_id).get()
    resolutions = Resolution.objects.filter(
        patient=patient, doctor=request.user).order_by('date')
    partesDoUtilizador = ParteDoUtilizador.objects.filter(
        participante=patient, resolution__doctor=request.user).order_by('data')
    parte = Part.objects.all()

    # r = Resolution.objects.filter(patient = patient, part__part__name = "MentHA-Risk")
    existing_risk = ParteDoUtilizador.objects.filter(
        participante=patient, part__name="MentHA-Risk")

    user = request.user
    # print(request.user.groups.all())

    # isto server para ver qual dos grupos dos users para ver permições e deixar aceder o que ao que
    user_risk = None
    user_tudo = None

    if request.user.groups.filter(name='Administrador').exists():
        user_tudo = request.user.groups.filter(name='Administrador')
    if request.user.groups.filter(name='Avaliador').exists():
        user_tudo = request.user.groups.filter(name='Avaliador')
    if request.user.groups.filter(name='Avaliador-Risk').exists():
        user_risk = request.user.groups.filter(name='Avaliador-Risk')

    risk_area = Area.objects.get(id=47)
    pergunta_risk = Question.objects.get(id=189)

    area_soc = Area.objects.get(pk=78)
    pergunta_soc = Question.objects.get(pk=187)

    c = []
    age = calculate_age(patient.nascimento)
    answered_list = []
    percentage_list = []
    report_list = []

    form = AppointmentForm(request.POST or None)

    if request.method == "POST":
        form = AppointmentForm(request.POST)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.participante = patient
            obj.save()

            resolution = Resolution(
                part=obj, patient=patient, doctor=request.user)
            resolution.initialize_statistics()
            resolution.save()

    for parteDoUtilizador in partesDoUtilizador:
        resolution = resolutions.filter(
            part=parteDoUtilizador, patient=patient, doctor=request.user).order_by('date')
        if not resolution:
            answered_list.append(0)
            percentage_list.append(0)
            report_list.append(None)
        else:
            s = resolution.get().statistics
            answered_list.append(s.get('total_answered'))
            percentage_list.append(s.get('total_percentage'))
            report_list.append(parteDoUtilizador.resolution.first().id)

    for cuidador in Cuidador.objects.all():
        if cuidador in patient.cuidadores.all():
            c.append(cuidador.nome)

    cuidadores = ", ".join(c)

    # Copia da patient_overview_view
    areas = Area.objects.order_by('part__order', 'order').all()
    # print(areas)
    overview_list = []

    # Getting all unique area names to display on html page
    for area in areas:
        for instrument in Instrument.objects.filter(area=area):
            text = area.name
            if instrument.name != "None":
                text = text + " - " + instrument.name
            overview_list.append(text)

    ow_l = dict.fromkeys(x for x in overview_list)

    percentages = {}
    # Getting % done per area, per part, to display done or not done on HTML
    for part_do_utilizador in partesDoUtilizador:
        percentages[part_do_utilizador.id] = {}
        part = part_do_utilizador.part
        r = resolutions.filter(part=part_do_utilizador)
        if len(r) < 1:
            for text in overview_list:
                this_area = Area.objects.filter(
                    part=part, name=text.split(' - ')[0]).order_by('order')
                if len(this_area) < 1:
                    percentages[part_do_utilizador.id][text] = "does not exist"
                else:
                    percentages[part_do_utilizador.id][text] = "not done"
        else:
            r = r.get()
            for text in overview_list:
                area_text = text.split(' - ')[0]
                this_area = Area.objects.filter(
                    part=part, name=area_text).order_by('order')
                if len(this_area) < 1:
                    percentages[part_do_utilizador.id][text] = "does not exist"
                else:
                    this_area = this_area.get()
                    instruments = Instrument.objects.filter(
                        area=this_area, area__part=part).order_by('order')
                    for instrument in instruments:
                        if instrument.name != 'None':
                            if f"{this_area.id}" in r.statistics and f"{instrument.id}" in r.statistics[f"{this_area.id}"]:
                                p = r.statistics[f"{this_area.id}"][f"{instrument.id}"].get(
                                    'percentage')
                            else:
                                p = 0

                        else:
                            p = r.statistics[f"{this_area.id}"].get(
                                'percentage')

                        if p == 100:
                            percentages[part_do_utilizador.id][text] = "done"
                        else:
                            percentages[part_do_utilizador.id][text] = "not done"

    context = {'patient': patient, 'cuidadores': cuidadores, 'resolutions': resolutions, 'partesdoutilizador': zip(partesDoUtilizador, answered_list, percentage_list, report_list),
               'partedoutilizador': partesDoUtilizador,
               'overview_list': ow_l,
               'percentages': percentages, 'age': age,
               'parte': parte,
               'risk_area': risk_area,
               'pergunta_risk': pergunta_risk,
               'form': form,
               'user': user,
               'user_risk': user_risk,
               'user_tudo': user_tudo,
               'existing_risk': existing_risk,
               'area_soc': area_soc,
               'pergunta_soc': pergunta_soc,
               }

    return render(request, 'protocolo/profile.html', context)


@login_required(login_url='login')
def gds_overview_view(request, protocol_id, part_id, area_id, instrument_id, patient_id):
    patient = Participante.objects.filter(pk=patient_id).get()
    part = ParteDoUtilizador.objects.get(pk=part_id)
    r = Resolution.objects.filter(
        patient=patient, doctor=request.user, part=part).get()
    answers = Answer.objects.filter(resolution=r)
    name_list = ['Respondente #1', 'Respondente #2', 'Respondente #3']
    overview_list = []
    questions = Question.objects.filter(section__dimension__instrument__name="GDS", name__in=name_list
                                        ).order_by('order')

    rowspan_estadio1 = len(
        questions[0].possible_answers.all().filter(description='1'))
    rowspan_estadio2 = len(
        questions[0].possible_answers.all().filter(description='2'))
    rowspan_estadio3 = len(
        questions[0].possible_answers.all().filter(description='3'))
    rowspan_estadio4 = len(questions[0].possible_answers.all().filter(description='4')) + len(
        questions[0].possible_answers.all().filter(description='4 - 2'))
    rowspan_estadio5 = len(questions[0].possible_answers.all().filter(description='5')) + len(
        questions[0].possible_answers.all().filter(description='5 - 2'))
    rowspan_estadio6 = len(questions[0].possible_answers.all().filter(description='6')) + len(
        questions[0].possible_answers.all().filter(description='6 - 2')) + len(
        questions[0].possible_answers.all().filter(description='6 - 3')) + len(
        questions[0].possible_answers.all().filter(description='6 - 4'))
    rowspan_estadio7 = len(
        questions[0].possible_answers.all().filter(description='7'))

    answers_dict = {'1': {},
                    '2': {},
                    '3': {},
                    }

    for q in questions:
        for pa in q.possible_answers.all():
            respondente = q.name.split('#')[1]
            a = Answer.objects.filter(question=q, resolution=r)
            if len(a) > 0:
                a = a.get()
                if a.MCCAnswer is not None:
                    for mca in a.MCCAnswer.all():
                        if mca.choice == pa:
                            answers_dict[respondente][pa.name] = "checked"

    rowspans = {
        '1': rowspan_estadio1,
        '2': rowspan_estadio2,
        '3': rowspan_estadio3,
        '4': rowspan_estadio4,
        '5': rowspan_estadio5,
        '6': rowspan_estadio6,
        '7': rowspan_estadio7,
    }

    list_2, list_3, list_4, list_5, list_6, list_7 = [], [], [], [], [], []

    for pa in questions[0].possible_answers.all():
        if pa.description == '2':
            list_2.append(pa.name)
        elif pa.description == '3':
            list_3.append(pa.name)
        elif pa.description == '4' or pa.description == "4 - 2":
            list_4.append(pa.name)
        elif pa.description == '5' or pa.description == "5 - 2":
            list_5.append(pa.name)
        elif pa.description == '6' or pa.description == "6 - 2" or pa.description == "6 - 3" or pa.description == "6 - 4":
            list_6.append(pa.name)
        elif pa.description == '7':
            list_7.append(pa.name)

    questions_dict = {'2': list_2,
                      '3': list_3,
                      '4': list_4,
                      '5': list_5,
                      '6': list_6,
                      '7': list_7,
                      }

    max_2, max_3, max_4, max_5, max_6, max_7 = [0, 0, 0], [
        0, 0, 0], [0, 0, 0], [0, 0, 0], [0, 0, 0], [0, 0, 0]
    for q in questions:
        respondente = q.name.split('#')[1]
        a = Answer.objects.filter(question=q, resolution=r)
        if len(a) > 0:
            a = a.get()
            for mca in a.MCCAnswer.all():
                if mca.choice.description == '2':
                    max_2[int(respondente) - 1] += 1
                elif mca.choice.description == '3':
                    max_3[int(respondente) - 1] += 1
                elif mca.choice.description in ['4', '4 - 2']:
                    max_4[int(respondente) - 1] += 1
                elif mca.choice.description in ['5', '5 - 2']:
                    max_5[int(respondente) - 1] += 1
                elif mca.choice.description in ['6', '6 - 2', '6 - 3', '6 - 4']:
                    max_6[int(respondente) - 1] += 1
                elif mca.choice.description in ['7']:
                    max_7[int(respondente) - 1] += 1
    total_2, total_3, total_4, total_5, total_6, total_7 = 0, 0, 0, 0, 0, 0
    for pa in questions[0].possible_answers.all():
        if pa.description == '2':
            total_2 += 1
        elif pa.description == '3':
            total_3 += 1
        elif pa.description in ['4', '4 - 2']:
            total_4 += 1
        elif pa.description in ['5', '5 - 2']:
            total_5 += 1
        elif pa.description in ['6', '6 - 2', '6 - 3', '6 - 4']:
            total_6 += 1
        elif pa.description in ['7']:
            total_7 += 1

    max_total_dict = {'2': [max(max_2), total_2],
                      '3': [max(max_3), total_3],
                      '4': [max(max_4), total_4],
                      '5': [max(max_5), total_5],
                      '6': [max(max_6), total_6],
                      '7': [max(max_7), total_7],
                      }

    context = {'questions': questions,
               'question': questions[0],
               'rowspans': rowspans,
               'dict': answers_dict,
               'questions_dict': questions_dict,
               'max_total_dict': max_total_dict,
               'protocol_id': protocol_id,
               'part_id': part_id,
               'area_id': area_id,
               'instrument_id': instrument_id,
               'patient_id': patient_id,
               }
    return render(request, 'protocolo/gds_overview.html', context)


@login_required(login_url='login')
def patient_overview_view(request, participant_id):
    patient = Participante.objects.filter(pk=participant_id).get()
    resolutions = Resolution.objects.filter(
        patient=patient).order_by('part__order')
    areas = Area.objects.order_by('part__order', 'order').all()
    # print(areas)
    parts = Part.objects.all()
    overview_list = []

    # Getting all unique area names to display on html page
    for area in areas:
        for instrument in Instrument.objects.filter(area=area):
            text = area.name
            if instrument.name != "None":
                text = text + " - " + instrument.name
            overview_list.append(text)

    ow_l = dict.fromkeys(x for x in overview_list)

    percentages = {1: {}, 2: {}, 3: {}, 4: {}, 5: {}, 6: {}}
    # Getting % done per area, per part, to display done or not done on HTML
    for part in parts:
        r = resolutions.filter(part=part)
        if len(r) < 1:
            for text in overview_list:
                this_area = Area.objects.filter(
                    part=part, name=text.split(' - ')[0]).order_by('order')
                if len(this_area) < 1:
                    percentages[part.order][text] = "does not exist"
                else:
                    percentages[part.order][text] = "not done"
        else:
            r = r.get()
            for text in overview_list:
                area_text = text.split(' - ')[0]
                this_area = Area.objects.filter(
                    part=part, name=area_text).order_by('order')
                if len(this_area) < 1:
                    percentages[part.order][text] = "does not exist"
                else:
                    this_area = this_area.get()
                    instruments = Instrument.objects.filter(
                        area=this_area, area__part=part).order_by('order')
                    for instrument in instruments:
                        if instrument.name != 'None':
                            if f"{this_area.id}" in r.statistics and f"{instrument.id}" in r.statistics[f"{this_area.id}"]:
                                p = r.statistics[f"{this_area.id}"][f"{instrument.id}"].get(
                                    'percentage')
                            else:
                                p = 0
                        else:
                            p = r.statistics[f"{this_area.id}"].get(
                                'percentage')

                        if p == 100:
                            percentages[part.order][text] = "done"
                        else:
                            percentages[part.order][text] = "not done"

    # print(percentages.get(2))
    # print(overview_list)
    context = {'patient': patient,
               'resolutions': resolutions,
               'parts': parts,
               'overview_list': ow_l,
               'percentages': percentages,
               }
    return render(request, 'protocolo/patient_overview.html', context)


# quero obter o json com as respostas de um paciente de risk
def risk_json(path, smoking, idade, colesterol, hipertensao):
    # abrir json risk_men
    data = open_json(path)
    for i in data:
        if (i == smoking):
            for j in data[i]:
                min = j.split('-')[0]
                max = j.split('-')[1]
                min = int(min)
                max = int(max)
                idade = int(idade)
                if (idade in range(min, max+1)):
                    for k in data[i][j]:
                        min = k.split('-')[0]
                        max = k.split('-')[1]
                        min = int(min)
                        max = int(max)
                        hipertensao = int(hipertensao)
                        if (hipertensao in range(min, max+1)):
                            for l in data[i][j][k]:
                                min = l.split('-')[0]
                                max = l.split('-')[1]
                                min = float(min)
                                max = float(max)
                                colesterol = float(colesterol)
                                if (colesterol in float_range(min, max+0.1)):
                                    return data[i][j][k][l]


def generate_id():
    # Generate a random number
    random_num = str(random.randint(0, 99999999)).encode()

    # Generate a SHA-256 hash of the random number.
    hash_obj = hashlib.sha256(random_num)
    hex_digit = hash_obj.hexdigest()

    return hex_digit[:10]
# funcao que abre um ficheiro Json e devolve um dicionario


def open_json(file_name):
    with open(file_name) as json_file:
        data = json.load(json_file)
        return data
# funcao que faça um float range de 0.1 em 0.1


def float_range(start, stop, step=0.1):
    while start < stop:
        yield round(start, 1)
        start += step


def gera_relatorio_risk_pdf(parte_risk, patient, username, genero, boolean_idade, boolean_pressao):

    document = Document()
    # conversao para boleans
    if parte_risk.diabetes == "True":
        parte_risk.diabetes = True
    else:
        parte_risk.diabetes = False
    if parte_risk.enfarte == "True":
        parte_risk.enfarte = True
    else:
        parte_risk.enfarte = False
    if parte_risk.doenca_rins == "True":
        parte_risk.doenca_rins = True
    else:
        parte_risk.doenca_rins = False
    if parte_risk.hipercolestrol == "True":
        parte_risk.hipercolestrol = True
    else:
        parte_risk.hipercolestrol = False
    if parte_risk.doenca_cognitiva == "True":
        parte_risk.doenca_cognitiva = True
    else:
        parte_risk.doenca_cognitiva = False
    if parte_risk.avc == "True":
        parte_risk.avc = True
    else:
        parte_risk.avc = False
    if parte_risk.doenca_pernas == "True":
        parte_risk.doenca_pernas = True
    else:
        parte_risk.doenca_pernas = False   
    if parte_risk.pre_diabetico == "True":
        parte_risk.pre_diabetico = True
    else:
        parte_risk.pre_diabetico = False

    # guardar valores para uso
    risco_baixo = ""
    hipercolesterol = ""
    drc = ""
    fuma = ""
    cardio = ""
    diabetes = ""

    # conversao para ints
    idade = int(parte_risk.idade)
    risco = int(parte_risk.risco_de_enfarte)
    # cores
    color_low = '00FF00'  # Green
    color_moderate = 'FFD700'  # Yellow 'FFFF00'
    color_high = 'FF0000'  # Red
    color_very_high = '8B0000'  # Dark Red

    # Cabeçalho
    paragraph = document.add_paragraph(f'MentHA-Risk')

    # para pôr em itálico (chato... talvez exista algo melhor)
    for run in paragraph.runs:
        run.font.italic = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_heading(
        f'Avaliação de risco Cardiovascular de {patient.__str__()}', 0)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Relatório

    # Dados Sócio-demográficos
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    run = paragraph.add_run('Dados Socio-demográficos')
    run.bold = True
    run.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Nome:')
    run.bold = True
    paragraph.add_run(f' {patient.__str__()}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Instituição de Referência:')
    run.bold = True
    paragraph.add_run(f' {patient.referenciacao}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Local de Residência:')
    run.bold = True
    paragraph.add_run(f' {patient.localizacao}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'E-mail:')
    run.bold = True
    paragraph.add_run(f' {patient.email}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Dados para SCORE-2
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    run = paragraph.add_run('Dados para SCORE-2')
    run.bold = True
    run.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph = document.add_paragraph()

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Idade:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.idade}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Sexo:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.sexo}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if parte_risk.fumador == 'smoking':
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'É fumador: ')
        run.bold = True
        paragraph.add_run(f'Sim')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        fuma = "É fumador"

    if  parte_risk.fumador == 'exSmoking':
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Já foi fumador: ')
        run.bold = True
        paragraph.add_run(f'Sim')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        fuma = "Já foi fumador"

    if  parte_risk.fumador == 'naoSeSabe':
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Fumador: ')
        run.bold = True
        paragraph.add_run(f'Não se tem conhecimento')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        fuma = "Não se sabe"

    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'É fumador:')
        run.bold = True
        paragraph.add_run(f' Não')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        fuma = "Não fuma"

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Pressão arterial sistólica:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.pressao_arterial} mmHg')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Dados Complementares
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    run = paragraph.add_run('Dados Complementares')
    run.bold = True
    run.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Peso:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.peso} kg')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Pat:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.pat_id}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'IMC:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.imc}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Altura:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.altura} cm')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Colesterol total:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.colestrol_total} mg/Dl')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Colesterol HDL:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.colestrol_hdl} mg/Dl')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Colesterol Não HDL:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.colestrol_nao_hdl} mg/Dl')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'TG- Triglicemia:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.tg} mg/Dl')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'LDL- Lipoproteínas de baixa densidade:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.ldl} mg/Dl')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'CHOL/HDL- Relação Chol por HDL:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.cholhdl} mg/Dl')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Batimentos cardiacos:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.batimentos} bpm')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    if parte_risk.diabetes:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Diabetes:')
        run.bold = True
        paragraph.add_run(f' Tem Diabetes')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        diabetes = "com Diabetes"
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Diabetes:')
        run.bold = True
        paragraph.add_run(f' Não tem Diabetes')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        diabetes = "sem Diabetes"
    if parte_risk.pre_diabetico:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Pré-Diabético?')
        run.bold = True
        paragraph.add_run(f' Sim, é pré-diabético')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Pré-Diabético?')
        run.bold = True
        paragraph.add_run(f' Não, não é pré-diabético')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if parte_risk.diabetes:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Há quantos anos o paciente tem diabétes:')
        run.bold = True
        paragraph.add_run(f' {parte_risk.anos_diabetes} anos')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        pass

    if parte_risk.avc:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            f'Teve algum AVC (Acidente Vascular Cerebral):')
        run.bold = True
        paragraph.add_run(f' Sim teve')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            f'Teve algum AVC (Acidente Vascular Cerebral):')
        run.bold = True
        paragraph.add_run(f' Não teve')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if parte_risk.enfarte:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Teve algum Enfarte (Ataque cardiaco):')
        run.bold = True
        paragraph.add_run(f' Sim teve')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        cardio = "com histório de evento cardiovascular major"
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Teve algum Enfarte (Ataque cardiaco):')
        run.bold = True
        paragraph.add_run(f' Não teve')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        cardio = "sem histório de evento cardiovascular major"
    if parte_risk.doenca_rins:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Tem doença dos Rins:')
        run.bold = True
        paragraph.add_run(f' Sim tem')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        drc = "e com doença renal crónica"
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Tem doença dos Rins:')
        run.bold = True
        paragraph.add_run(f' Não tem')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        drc = "sem doença renal crónica"
    if parte_risk.doenca_pernas:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Tem doença das Pernas:')
        run.bold = True
        paragraph.add_run(f' Sim tem')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Tem doença das Pernas:')
        run.bold = True
        paragraph.add_run(f' Não tem')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if parte_risk.hipercolestrol:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Historial de hipercolesterolemia familiar:')
        run.bold = True
        paragraph = document.add_paragraph(
            f' Existe familiares com hipercolesterolémia')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        hipercolesterol = "com hipercolesterolémia familiar"
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Historial de hipercolesterolemia familiar:')
        run.bold = True
        paragraph = document.add_paragraph(
            f' Não existe familiares com hipercolesterolémia')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        hipercolesterol = "sem hipercolesterolémia familiar"
    if parte_risk.doenca_cognitiva:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Tem doença cognitiva?')
        run.bold = True
        paragraph = document.add_paragraph(
            f' Sim tem doença cognitiva')
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Historial de hipercolesterolemia familiar:')
        run.bold = True
        paragraph = document.add_paragraph(
            f'Não tem doença cognitiva')
    
    # Resumo dos dados/Resultados
    #teste dados novos Intervalo de medição HBA1:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Intervalo de medição HBA1:')

    run.bold = True
    run.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Eag 68-356:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.eag_hba1} mg/Dl')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'IFCC: 20-130 (mmol/mol)')
    run.bold = True
    paragraph.add_run(f' {parte_risk.ifcc_hba1} mg/Dl')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'NGSP: 4-14 (%)')
    run.bold = True
    paragraph.add_run(f' {parte_risk.ngsp_hba1} mg/Dl')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Resultados')
    run.bold = True
    run.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if boolean_pressao and boolean_idade:

        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            f'Probabilidade de ter um Risco Cardiovascular:')
        run.bold = True
        paragraph.add_run(f' {parte_risk.risco_de_enfarte} %')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        if ((risco <= 2 and idade < 50) or (risco < 5 and (idade >= 50 and idade <= 69)) or (risco <= 5 and idade >= 70)):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f'Classificação do Risco Cardiovascular:')
            run.bold = True
            paragraph.add_run(f' Baixo')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            add_cores(paragraph, color_low)
            risco_baixo = "Baixo"
        elif (((risco > 2 or risco <= 8) and idade < 50) or (risco < 5 and (idade >= 50 or idade <= 69)) or (risco <= 5 and idade >= 70)):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f'Classificação do Risco Cardiovascular:')
            run.bold = True
            paragraph.add_run(f' Moderado')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            add_cores(paragraph, color_moderate)
            risco_baixo = "Moderado"
        elif (((risco >= 2 and risco < 8) and idade < 50) or ((risco >= 5 and risco < 10) and (idade >= 50 and idade <= 69)) or ((risco >= 7 and risco < 15) and idade >= 70)):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f'Classificação do Risco Cardiovascular:')
            run.bold = True
            paragraph.add_run(f' Alto')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            add_cores(paragraph, color_high)
            risco_baixo = "Alto"
        elif ((risco >= 7 and idade < 50) or (risco >= 10 and (idade >= 50 and idade <= 69)) or (risco >= 15 and idade >= 70)):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f'Classificação do Risco Cardiovascular:')
            run.bold = True
            paragraph.add_run(f' Muito Alto')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            add_cores(paragraph, color_very_high)
            risco_baixo = "Elevado"

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Descrição: ')
        run.bold = True
        if parte_risk.sexo == 'Masculino':
            paragraph.add_run(f'Um Homen de {idade} anos {cardio},{diabetes} ,{hipercolesterol} e {drc}. O valor de colesterol total de {parte_risk.colestrol_total} mg/dl e o colesterol HDL de {parte_risk.colestrol_hdl} mg/dl; logo o colesterol não HDL de {parte_risk.colestrol_nao_hdl} mg/dl. {fuma} e o valor da pressão arterial sistolica é de {parte_risk.pressao_arterial} mmHg. O score de risco é de {parte_risk.risco_de_enfarte}% o que significa que tem {risco_baixo} risco cardiovascular.')
        else:
            paragraph.add_run(f'Uma Mulher de {idade} anos {cardio},{diabetes} ,{hipercolesterol} e {drc}. O valor de colesterol total de {parte_risk.colestrol_total} mg/dl e o colesterol HDL de {parte_risk.colestrol_hdl} mg/dl; logo o colesterol não HDL de {parte_risk.colestrol_nao_hdl} mg/dl. {fuma} e o valor da pressão arterial sistolica é de {parte_risk.pressao_arterial} mmHg. O score de risco é de {parte_risk.risco_de_enfarte}% o que significa que tem {risco_baixo} risco cardiovascular.')

        if idade < 70 and idade >= 40:

            # nome_ficheiro_imagem = 'SCORE-2-1-' +patient.__str__()+generate_id()+'.png'
            img_path = os.path.join(
                os.getcwd(), 'protocolo', 'static', 'protocolo', 'img', 'SCORE-2-1.png')
            # new_img_path = os.path.join(os.getcwd(), 'protocolo\static\protocolo\img\img-report\SCORE-2-1-'+patient.__str__()+generate_id()+'.png')

            document.add_picture(
                img_path, width=Inches(4.5), height=Inches(4.5))
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        elif idade >= 70 and idade < 90:

            # nome_ficheiro_imagem = 'SCORE-2-1-' +patient.__str__()+generate_id()+'.png'
            img_path = os.path.join(
                os.getcwd(), 'protocolo', 'static', 'protocolo', 'img', 'SCORE-2-90.png')

            document.add_picture(
                img_path, width=Inches(4.5), height=Inches(4.5))
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # fazer parte dinamica para mostrar os valores de risco
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            f'NÃO EXISTE DADOS SUFICIENTES PARA CALCULAR O RISCO DE TER UM ACIDENTE CARDIOVASCULAR')
        run.bold = True
        run.font.size = Pt(14)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            f'NOTA: É necessário ter uma idade superior a 40 e menor que 90 anos e uma pressão arterial sistólica superior a 100 mmHg e menos de 179mmHg')
        run.bold = True
        run.font.size = Pt(14)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Recomendações
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Recomendações')
    run.bold = True
    run.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'O Paciente deve:')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'  •	Deixar de fumar')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'  •	Reduzir o consumo de álcool')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'  •	Reduzir o consumo de sal')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        f'  •	Reduzir o consumo de gorduras saturadas e colesterol')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'  •	Reduzir o consumo de açúcares e doces')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        f'  •	Fazer uma alimentação saudável e equilibrada')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'  •	Praticar exercício físico regularmente')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Comentário do Avaliador
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Comentário do Avaliador:')
    run.bold = True
    run.font.size = Pt(14)
    paragraph = document.add_paragraph(f'{parte_risk.comentario}')
    run = paragraph.add_run(f'Recomendações adicionais:')
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'{parte_risk.recomendacoes}')
    run.bold = True
    run.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Add the image
    # img_path = os.path.join(os.getcwd(), 'img', 'example.jpg')
    # document.add_picture(img_path, width=Inches(3), height=Inches(3))asdsad

    # Assinatura
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph(f'O avaliador, {username}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph = document.add_paragraph()

    # imagens logos no fundo do documento
    # cwd = os.getcwd()
    # cwd2 = os.path.join(cwd, 'mentha', 'static', 'img', 'img-logo','ulht.png')
    section = document.sections[0]

    # Configurar o rodapé
    footer = section.footer
    # Certifique-se de que o rodapé não esteja vinculado ao anterior
    footer.is_linked_to_previous = False

# Criar parágrafo vazio no rodapé para adicionar as imagens
    paragraph = footer.paragraphs[0]
    image_paths = [
        os.path.join(os.getcwd(), 'protocolo', 'static',
                     'protocolo', 'img', 'img-logo', 'ulht.png'),
        os.path.join(os.getcwd(), 'protocolo', 'static',
                     'protocolo', 'img', 'img-logo', 'dgs_footer.png'),
        os.path.join(os.getcwd(), 'protocolo', 'static',
                     'protocolo', 'img', 'img-logo', 'adebe.png'),
        os.path.join(os.getcwd(), 'protocolo', 'static',
                     'protocolo', 'img', 'img-logo', 'copelabs.jpg'),
        os.path.join(os.getcwd(), 'protocolo', 'static',
                     'protocolo', 'img', 'img-logo', 'gira.png'),
        os.path.join(os.getcwd(), 'protocolo', 'static',
                     'protocolo', 'img', 'img-logo', 'cvp.jpg'),
        os.path.join(os.getcwd(), 'protocolo', 'static',
                     'protocolo', 'img', 'img-logo', 'familiarmente.jpg'),
        os.path.join(os.getcwd(), 'protocolo', 'static',
                     'protocolo', 'img', 'img-logo', 'elo.png'),
        os.path.join(os.getcwd(), 'protocolo', 'static',
                     'protocolo', 'img', 'img-logo', 'mentha-logo.png')
    ]
    for i, image_path in enumerate(image_paths):
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(0.3), height=Inches(0.3))
        if i < len(image_path)-1:
            run.add_text(" ")

    # img_path6 = os.path.join(os.getcwd(), 'mentha\\static\mentha\\pareceiros_sm\\dgs_footer.png')
    # document.add_picture(img_path6, width=Inches(1.5), height=Inches(1.5))
    # img_path = os.path.join(os.getcwd(), 'protocolo\static\protocolo\img\logo4.png')
    # document.add_picture(img_path, width=Inches(1.5), height=Inches(1.5))
    # img_path = os.path.join(os.getcwd(), 'protocolo\static\protocolo\img\logo5.png')
    # document.add_picture(img_path, width=Inches(1.5), height=Inches(1.5))

    # Save the Word document
    nome_ficheiro = 'Risco_Cardiovascular' + \
        '_' + patient.__str__() + generate_id()
    nome_ficheiro = nome_ficheiro.replace(" ", "")
    nome_ficheiro = nome_ficheiro.replace("/", "_")
    docx_path = os.path.join(os.getcwd(), f'{nome_ficheiro}.docx')
    print('Saving', docx_path)
    document.save(docx_path)

    # Create a Django File object from the Word document
    with open(docx_path, 'rb') as f:
        docx_data = io.BytesIO(f.read())

    parte_risk.relatorio_word.save(f'{nome_ficheiro}.docx', docx_data)
    parte_risk.save()

    os.remove(docx_path)


# (resolution, chc, coop, rel,report,instruments,numero_abvd)
def gera_relatorio_parte(resolution, chc, coop, rel, reporte, report):
    # chc vai ser uma lista de listas

    document = Document()
    # Cabeçalho
    paragraph = document.add_paragraph(
        f'Relatório do Protocolo de Avaliação MentHA')

    # para pôr em itálico (chato... talvez exista algo melhor)
    for run in paragraph.runs:
        run.font.italic = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_heading(
        f'{resolution.part.part.name} de {resolution.pessoa.info_sensivel.nome}', 0)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph = document.add_paragraph()

    # Tabela
    # nr_areas = resolution.part.part.number_of_areas
    areas = resolution.part.part.area

    instrumentes = None
    dimensoes = None
    # table = document.add_table(rows=1, cols=3, style="Table Grid")

    # # Cabeçalho da Tabela
    # heading_row = table.rows[0].cells
    # heading_row[0].text = "Área de Avaliação"
    # heading_row[1].text = "Resultado"

    for area in areas:
        name = area.name

        # add new row to table
        # data_row = table.add_row().cells

        # add headings
        if name == "Questionário Sociodemográfico":
            continue

        # data_row[0].text = area.name

        if name == 'Consciência, Humor e Comportamento':

            paragraph = document.add_paragraph()
            run = paragraph.add_run('Consciência, Humor e Comportamento')
            run.bold = True
            run.font.size = Pt(16)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph = document.add_paragraph()
            chc_consciencia = ", ".join([str(chc[0])]if len(chc) > 0 else [])

            # chc_consciencia.replace("['", "").replace(","," | ").replace("']", "")

            chc_motora = ", ".join([str(chc[1])]if len(chc) > 0 else [])
            # chc_motora.replace("['", "").replace(","," | ").replace("']", "")
            chc_humor = ", ".join([str(chc[2])]if len(chc) > 0 else [])
            # chc_humor.replace("['", "").replace(","," | ").replace("']", "")

            if len(chc_consciencia) > 0:
                runs = paragraph.add_run('Consciência: ')
                runs.bold = True
                paragraph.add_run(chc_consciencia)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                paragraph = document.add_paragraph()

            if len(chc_motora) > 0:
                run = paragraph.add_run('Atividade Motora: ')
                run.bold = True
                paragraph.add_run(chc_motora)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                paragraph = document.add_paragraph()

            if len(chc_humor) > 0:
                run = paragraph.add_run('Humor: ')
                run.bold = True
                paragraph.add_run(chc_humor)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                paragraph = document.add_paragraph()

        if name == 'Cooperação dada na entrevista':
            if len(coop) > 0:
                run = paragraph.add_run('Cooperação dada na entrevista')
                run.bold = True
                run.font.size = Pt(16)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph = document.add_paragraph()
                run = paragraph.add_run('Cooperação dada na Entrevista: ')
                run.bold = True
                paragraph.add_run(coop)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                paragraph = document.add_paragraph()

        if name == 'Relação com o Avaliador':
            if len(rel) > 0:
                run = paragraph.add_run('Relação com o Avaliador')
                run.bold = True
                run.font.size = Pt(16)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph = document.add_paragraph()
                run = paragraph.add_run('Relação com o Avaliador: ')
                run.bold = True
                paragraph.add_run(rel)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                paragraph = document.add_paragraph()

        if name == 'Funcionalidade':
            names = []
            quotation = []
            instrumentes = Instrument.objects.filter(area=area)
            paragraph = document.add_paragraph()
            paragraph = document.add_paragraph()
            run = paragraph.add_run('Funcionalidade')
            run.bold = True
            run.font.size = Pt(16)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph = document.add_paragraph()
            for instrument in instrumentes:
                if instrument.name != None:
                    if 'ABVD' in instrument.name:
                        paragraph = document.add_paragraph()
                        run = paragraph.add_run('ABVD')
                        run.bold = True
                        run.font.size = Pt(12)
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()
                        dimensoes = Dimension.objects.filter(
                            instrument=instrument)
                        reporte = Report.objects.get(resolution=resolution)
                        for dimensao in dimensoes:
                            if dimensao.name == 'Atividades Corporais':
                                names.append(dimensao.name)
                                corporal = reporte.abvd_atividades_corporais_quotation
                                quotation.append(corporal)
                                run = paragraph.add_run(
                                    'Atividades Corporais:')
                                run.bold = True
                                paragraph.add_run(f' {corporal}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                            if dimensao.name == 'Atividades Motoras':
                                names.append(dimensao.name)
                                motoras = reporte.abvd_atividades_motoras_quotation
                                quotation.append(motoras)
                                run = paragraph.add_run('Atividades Motoras:')
                                run.bold = True
                                paragraph.add_run(f' {motoras}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()
                            if dimensao.name == 'Atividades Sensoriais':
                                names.append(dimensao.name)
                                sensoriais = reporte.abvd_atividades_sensoriais_quotation
                                quotation.append(sensoriais)
                                run = paragraph.add_run(
                                    'Atividades Sensoriais:')
                                run.bold = True
                                paragraph.add_run(f' {sensoriais}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()
                            if dimensao.name == 'Atividades Mentais':
                                names.append(dimensao.name)
                                mentais = reporte.abvd_atividades_mentais_quotation
                                quotation.append(mentais)
                                run = paragraph.add_run('Atividades Mentais')
                                run.bold = True
                                paragraph.add_run(f' {mentais}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()
                        grau = reporte.abvd_evaluation
                        run = paragraph.add_run('Grau de Dependência:')
                        run.bold = True
                        paragraph.add_run(f' {grau}')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        if instrument.name in report[name] and report[name][instrument.name]['img']:
                            grafico = report[name][instrument.name]['img']
                            document.add_picture(grafico, width=Inches(5))
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            pass

                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        paragraph = document.add_paragraph()

                    if 'AIVD' in instrument.name:
                        paragraph = document.add_paragraph()
                        run = paragraph.add_run('AIVD')
                        run.bold = True
                        run.font.size = Pt(12)
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()
                        dimensoes = Dimension.objects.filter(
                            instrument=instrument)
                        reporte = Report.objects.get(resolution=resolution)
                        for dimensao in dimensoes:
                            if dimensao.name == 'None':
                                grau = reporte.aivd_evaluation
                                run = paragraph.add_run('Grau de Dependência:')
                                run.bold = True
                                paragraph.add_run(f' {grau}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

        if name == 'Psicossintomatologia':
            paragraph = document.add_paragraph()
            run = paragraph.add_run('Psicossintomatologia')
            run.bold = True
            run.font.size = Pt(16)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph = document.add_paragraph()
            instrumentes = Instrument.objects.filter(area=area)
            for instrument in instrumentes:
                if instrument.name != None:
                    if 'BSI' in instrument.name:
                        paragraph = document.add_paragraph()
                        run = paragraph.add_run('BSI')
                        run.bold = True
                        run.font.size = Pt(12)
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()
                        dimensoes = Dimension.objects.filter(
                            instrument=instrument)
                        reportar = Report.objects.get(resolution=resolution)

                        for dimensao in dimensoes:
                            if dimensao == None:
                                paragraph.add_run('Somatização')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()
                            else:
                                somatizacao = reportar.bsi_somatizacao
                                run = paragraph.add_run('Somatização: ')
                                run.bold = True
                                paragraph.add_run(f'{somatizacao}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                                obsessao = reportar.bsi_obssessivo_compulsivo
                                run = paragraph.add_run(
                                    'Obsessões-Compulsões: ')
                                run.bold = True
                                paragraph.add_run(f'{obsessao}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                                depressao = reportar.bsi_depressao
                                run = paragraph.add_run('Depressão: ')
                                run.bold = True
                                paragraph.add_run(f'{depressao}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                                sensibilidade = reportar.bsi_sensibilidade_interpessoal
                                run = paragraph.add_run(
                                    'Sensibilidade Interpessoal: ')
                                run.bold = True
                                paragraph.add_run(f'{sensibilidade}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                                ansiadade = reportar.bsi_ansiedade
                                run = paragraph.add_run('Ansiedade: ')
                                run.bold = True
                                paragraph.add_run(f'{ansiadade}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                                hostilidade = reportar.bsi_hostilidade
                                run = paragraph.add_run('Hostilidade: ')
                                run.bold = True
                                paragraph.add_run(f'{hostilidade}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                                fobia = reportar.bsi_ansiedade_fobica
                                run = paragraph.add_run('Ansiedade Fóbica: ')
                                run.bold = True
                                paragraph.add_run(f'{fobia}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                                paranoide = reportar.bsi_paranoide
                                run = paragraph.add_run('Ideação Paranoide: ')
                                run.bold = True
                                paragraph.add_run(f'{paranoide}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                                psicoticismo = reportar.bsi_psicoticismo
                                run = paragraph.add_run('Psicoticismo: ')
                                run.bold = True
                                paragraph.add_run(f'{psicoticismo}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                                igs = reportar.bsi_igs
                                run = paragraph.add_run(
                                    'Índice Global de Sintomatologia: ')
                                run.bold = True
                                paragraph.add_run(f'{igs}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                                tsp = reportar.bsi_tsp
                                run = paragraph.add_run(
                                    'Total de Sintomas Positivos: ')
                                run.bold = True
                                paragraph.add_run(f'{tsp}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                                isp = reportar.bsi_isp
                                run = paragraph.add_run(
                                    'Índice de Sintomas Positivos: ')
                                run.bold = True
                                paragraph.add_run(f'{isp}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                        if instrument.name in report[name] and report[name][instrument.name]['img']:
                            grafico = report[name][instrument.name]['img']
                            document.add_picture(grafico, width=Inches(5))
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            pass

        if name == 'Cognição':
            paragraph = document.add_paragraph()
            run = paragraph.add_run('Cognição')
            run.bold = True
            run.font.size = Pt(16)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph = document.add_paragraph()
            instrumentes = Instrument.objects.filter(area=area)
            for instrument in instrumentes:
                if instrument.name != None:
                    if 'ACE-R' == instrument.name:
                        paragraph = document.add_paragraph()
                        run = paragraph.add_run('ACE-R')
                        run.bold = True
                        run.font.size = Pt(12)
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()
                        dimensoes = Dimension.objects.filter(
                            instrument=instrument)
                        reportar = Report.objects.get(resolution=resolution)
                        for dimensao in dimensoes:
                            if dimensao == 'Atenção e Orientação':
                                temporal = reportar.acer_atencao_orientacao_quotation
                                run = paragraph.add_run(
                                    'Atenção e Orientação: ')
                                run.bold = True
                                paragraph.add_run(f'{temporal}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                            if dimensao == 'Memória':
                                memoria = reportar.acer_memoria_quotation
                                run = paragraph.add_run('Memória: ')
                                run.bold = True
                                paragraph.add_run(f'{memoria}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                            if dimensao == 'Fluência':
                                fluencia = reportar.acer_fluencia_quotation
                                run = paragraph.add_run('Fluência: ')
                                run.bold = True
                                paragraph.add_run(f'{fluencia}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                            if dimensao == 'Linguagem':
                                linguagem = reportar.acer_linguagem_quotation
                                run = paragraph.add_run('Linguagem: ')
                                run.bold = True
                                paragraph.add_run(f'{linguagem}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()
                            if dimensao == 'Visuo-Espacial':
                                visuo = reportar.acer_visuo_espacial_quotation
                                run = paragraph.add_run('Visuo-Espacial: ')
                                run.bold = True
                                paragraph.add_run(f'{visuo}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                        grau_ = reportar.acer_evaluation
                        run = paragraph.add_run('Grau de Dependência: ')
                        run.bold = True
                        paragraph.add_run(f'{grau_}')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()
                        if instrument.name in report[name] and report[name][instrument.name]['img']:
                            grafico = report[name][instrument.name]['img']
                            document.add_picture(grafico, width=Inches(5))
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            pass

                    if instrument.name == 'MMSE':
                        paragraph = document.add_paragraph()
                        run = paragraph.add_run('MMSE')
                        run.bold = True
                        run.font.size = Pt(12)
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()
                        dimensoes = Dimension.objects.filter(
                            instrument=instrument)
                        reporte = Report.objects.get(resolution=resolution)
                        for dimensao in dimensoes:
                            if dimensao == 'Atenção e Orientação':
                                temporal = reporte.mmse_atencao_orientacao_quotation
                                run = paragraph.add_run(
                                    'Atenção e Orientação: ')
                                run.bold = True
                                paragraph.add_run(f'{temporal}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                            if dimensao == 'Memória':
                                memoria = reporte.mmse_memoria_quotation
                                run = paragraph.add_run('Memória: ')
                                run.bold = True
                                paragraph.add_run(f'{memoria}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                            if dimensao == 'Linguagem':
                                linguagem = reporte.mmse_linguagem_quotation
                                run = paragraph.add_run('Linguagem: ')
                                run.bold = True
                                paragraph.add_run(f'{linguagem}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                            if dimensao == 'Visuo-Espacial':
                                visuo = reporte.mmse_visuo_espacial_quotation
                                run = paragraph.add_run('Visuo-Espacial: ')
                                run.bold = True
                                paragraph.add_run(f'{visuo}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                                # for sesao in sesoes:
                                #     if sesao == 'Capacidade Visuo-Espacial (Pentágonos)':

                                #         continue
                        grau_ = reporte.mmse_evaluation
                        run = paragraph.add_run('Grau de Dependência: ')
                        run.bold = True
                        paragraph.add_run(f'{grau_}')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()
                        if instrument.name in report[name] and report[name][instrument.name]['img']:
                            grafico = report[name][instrument.name]['img']
                            document.add_picture(grafico, width=Inches(4))
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            pass
                    if instrument.name == 'Área Complementar':
                        dimensoes = Dimension.objects.filter(
                            instrument=instrument)
                        reporte = Report.objects.get(resolution=resolution)
                        for dimensao in dimensoes:
                            if dimensao.name == 'Memória Visual Imediata':
                                memoria_visual_imediata = reporte.ac_memoria_visual_imediata_quotation
                                run = paragraph.add_run(
                                    'Memória Visual Imediata: ')
                                run.bold = True
                                paragraph.add_run(f'{memoria_visual_imediata}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()
                            if dimensao.name == 'Memória Visual Diferida':
                                memoria_visual_diferida = reporte.ac_memoria_diferida_quotation
                                run = paragraph.add_run(
                                    'Memória Visual Diferida: ')
                                run.bold = True
                                paragraph.add_run(f'{memoria_visual_diferida}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()
                            if dimensao.name == 'Atenção Mantida':
                                atencao_mantida = reporte.ac_tmt_a_quotation
                                run = paragraph.add_run('Atenção Mantida: ')
                                run.bold = True
                                paragraph.add_run(f'{atencao_mantida}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

                            if dimensao.name == 'Atenção Dividida':
                                atencao_divida = reporte.ac_tmt_b_quotation
                                run = paragraph.add_run('Atenção Dividida: ')
                                run.bold = True
                                paragraph.add_run(f'{atencao_divida}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()
                            if dimensao.name == 'Orientação Esquerda/Direita':
                                orientacao = reporte.ac_gnosias_quotation
                                run = paragraph.add_run(
                                    'Orientação Esquerda/Direita: ')
                                run.bold = True
                                paragraph.add_run(f'{orientacao}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()
                            if dimensao.name == 'Abstração Verbal':
                                abstracao = reporte.ac_proverbios_quotation
                                run = paragraph.add_run('Abstração Verbal: ')
                                run.bold = True
                                paragraph.add_run(f'{abstracao}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()
                            if dimensao.name == 'Compreensão das Instruções':
                                instrucoes = reporte.ac_token_test_quotation
                                run = paragraph.add_run(
                                    'Compreensão das Instruções: ')
                                run.bold = True
                                paragraph.add_run(f'{instrucoes}')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                paragraph = document.add_paragraph()

        if name == 'Estado de Ânimo':
            instrumentes = Instrument.objects.filter(area=area)
            paragraph = document.add_paragraph()
            run = paragraph.add_run('Estado de Ânimo')
            run.bold = True
            run.font.size = Pt(16)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph = document.add_paragraph()

            for instrument in instrumentes:
                if instrument.name != None:
                    if instrument.name == 'PANAS':
                        paragraph = document.add_paragraph()
                        run = paragraph.add_run('PANAS')
                        run.bold = True
                        run.font.size = Pt(12)
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()
                        reporte = Report.objects.get(resolution=resolution)
                        interesado = reporte.panas_interessado
                        run = paragraph.add_run('Interessado: ')
                        run.bold = True
                        paragraph.add_run(f'{interesado}')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()

                        nervoso = reporte.panas_nervoso
                        run = paragraph.add_run('Nervoso: ')
                        run.bold = True
                        paragraph.add_run(f'{nervoso}')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()

                        entusiasmado = reporte.panas_entusiasmado
                        run = paragraph.add_run('Entusiasmado: ')
                        run.bold = True
                        paragraph.add_run(f'{entusiasmado}')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()

                        amedrontado = reporte.panas_amedrontado
                        run = paragraph.add_run('Amedrontado: ')
                        run.bold = True
                        paragraph.add_run(f'{amedrontado}')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()

                        inspirado = reporte.panas_inspirado
                        run = paragraph.add_run('Inspirado: ')
                        run.bold = True
                        paragraph.add_run(f'{inspirado}')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()

                        ativo = reporte.panas_ativo
                        run = paragraph.add_run('Ativo: ')
                        run.bold = True
                        paragraph.add_run(f'{ativo}')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()

                        assustado = reporte.panas_assustado
                        run = paragraph.add_run('Assustado: ')
                        run.bold = True
                        paragraph.add_run(f'{assustado}')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()

                        culpado = reporte.panas_culpado
                        run = paragraph.add_run('Culpado: ')
                        run.bold = True
                        paragraph.add_run(f'{culpado}')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()

                        determinado = reporte.panas_determinado
                        run = paragraph.add_run('Determinado: ')
                        run.bold = True
                        paragraph.add_run(f'{determinado}')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()

                        atormentado = reporte.panas_atormentado
                        run = paragraph.add_run('Atormentado: ')
                        run.bold = True
                        paragraph.add_run(f'{atormentado}')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()
        if name == 'Ansiedade e Depressão':
            instrumentes = Instrument.objects.filter(area=area)
            paragraph = document.add_paragraph()
            run = paragraph.add_run('Ansiedade e Depressão')
            run.bold = True
            run.font.size = Pt(16)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph = document.add_paragraph()
            for instrument in instrumentes:
                if instrument.name != None:
                    if instrument.name == 'HADS':
                        paragraph = document.add_paragraph()
                        run = paragraph.add_run('HADS')
                        run.bold = True
                        run.font.size = Pt(12)
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()
                        reporte = Report.objects.get(resolution=resolution)
                        ansiadade_evaluation = reporte.hads_estado_ansiedade_evaluation
                        ansiadade_quotation = reporte.hads_estado_ansiedade_quotation
                        run = paragraph.add_run('Estado de Ansiedade: ')
                        run.bold = True
                        paragraph.add_run(
                            f'{ansiadade_evaluation}({ansiadade_quotation})')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()

                        depressao_quotation = reporte.hads_estado_depressao_quotation
                        depressao_evaluation = reporte.hads_estado_depressao_evaluation
                        run = paragraph.add_run('Estado de Depressão: ')
                        run.bold = True
                        paragraph.add_run(
                            f'{depressao_evaluation}({depressao_quotation})')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()
        if name == 'Estadio de Deterioração':
            instrumentes = Instrument.objects.filter(area=area)
            paragraph = document.add_paragraph()
            run = paragraph.add_run('Estadio de Deterioração')
            run.bold = True
            run.font.size = Pt(16)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph = document.add_paragraph()
            reporte = Report.objects.get(resolution=resolution)
            for instrument in instrumentes:
                if instrument.name != None:
                    if instrument.name == 'GDS':
                        paragraph = document.add_paragraph()
                        run = paragraph.add_run('GDS')
                        run.bold = True
                        run.font.size = Pt(12)
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()

                        gds = reporte.gds_nivel
                        run = paragraph.add_run('Estadio de Deterioração: ')
                        run.bold = True
                        paragraph.add_run(f'{gds}')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph = document.add_paragraph()
        # data_row[1].text = "Texto"

    # Assinatura
    paragraph = document.add_paragraph(
        f'O avaliador, {resolution.doctor.username}')

    # Save the Word document
    nome_ficheiro = 'Avaliação_Psicológica' + '_' + \
        resolution.pessoa.info_sensivel.nome + generate_id()
    nome_ficheiro = nome_ficheiro.replace(" ", "")
    docx_path = os.path.join(os.getcwd(), f'{nome_ficheiro}.docx')
    print('Saving', docx_path)
    document.save(docx_path)

    # Create a Django File object from the Word document
    with open(docx_path, 'rb') as f:
        docx_data = io.BytesIO(f.read())

    # Create a Django File object from the PDF file

    reporte.word.save(f'{nome_ficheiro}.docx', docx_data)
    reporte.save()
    # Delete the temporary files
    os.remove(docx_path)
    # os.remove(pdf_path)


# funcao para calcular o IMC
def calcular_imc(peso, altura):
    # altura em metros
    # peso em kg
    # peso = float(peso)
    # altura = int(altura)
    # imc = peso/(altura*altura)
    imc = 1
    return imc
# funcao para por cores no word


def add_cores(paragraph, color):
    run = paragraph.add_run()
    run.text = u'\u25CF'  # Circle character
    run.font.color.rgb = RGBColor.from_string(color)


@login_required(login_url='diario:login')
def participante_update(request, participante_id):
    participante = Participante.objects.get(pk=participante_id)
    formParticipante = ParticipanteForm(
        request.POST or None, instance=participante)

    if request.method == 'POST':
        print("POST data: ", request.POST)  # Print POST data
        new_post = remove_empty_keys_from_post(request.POST)
        print("Processed POST data: ", new_post)  # Print processed POST data

        nome = new_post.get('nome', None)
        if nome is not None:
            # nome wont change, how can I make it?


            participante.info_sensivel.nome = nome
            participante.info_sensivel.save()
            print("nome after processing: ", participante.info_sensivel.nome)

        email = new_post.get('email', None)
        if email is not None:
            participante.info_sensivel.email = email
            participante.info_sensivel.save()
            
        diagnosticoPrincipal = new_post.get('diagnosticoPrincipal', None)
        if diagnosticoPrincipal is not None:
            participante.diagnosticoPrincipal = diagnosticoPrincipal
            participante.save()
            # Diagnostico Principal
        #Comorbilidades
        comorbilidades = new_post.get('comorbilidades', None)
        if comorbilidades is not None:
            participante.comorbilidades = comorbilidades
            participante.save()
        diagnosticoPrincipal = new_post.get('diagnosticoPrincipal', None)
        if diagnosticoPrincipal is not None:
            participante.diagnosticoPrincipal = diagnosticoPrincipal
            participante.save()
        

        telemovel = new_post.get('telemovel', None)
        if telemovel is not None:
            participante.info_sensivel.telemovel = telemovel
            participante.info_sensivel.save()

        imagem = request.FILES.get('imagem', None)
        if imagem is not None:
            participante.imagem = imagem

        escolaridade = new_post.get('escolaridade', None)
        if escolaridade is not None:
            participante.escolaridade = escolaridade

        nascimento = new_post.get('nascimento', None)
        if nascimento is not None:
            participante.nascimento = nascimento
            

        nacionalidade = new_post.get('nacionalidade', None)
        if nacionalidade is not None:
            participante.nacionalidade = nacionalidade

        sexo = new_post.get('sexo', None)
        if sexo is not None:
            participante.sexo = sexo

        localizacao = new_post.get('localizacao', None)
        if localizacao is not None:
            participante.localizacao = localizacao

        situacaoLaboral = new_post.get('situacaoLaboral', None)
        if situacaoLaboral is not None:
            participante.situacaoLaboral = situacaoLaboral

        profissaoPrincipal = new_post.get('profissaoPrincipal', None)
        if profissaoPrincipal is not None:
            participante.profissaoPrincipal = profissaoPrincipal

        situacaoEconomica = new_post.get('situacaoEconomica', None)
        if situacaoEconomica is not None:
            participante.situacaoEconomica = situacaoEconomica

        estadoCivil = new_post.get('estadoCivil', None)
        if estadoCivil is not None:
            participante.estadoCivil = estadoCivil

        agregadoFamiliar = new_post.get('agregadoFamiliar', None)
        if agregadoFamiliar is not None:
            participante.agregadoFamiliar = agregadoFamiliar

        autoAvaliacaoEstadoSaude = new_post.get(
            'autoAvaliacaoEstadoSaude', None)
        if autoAvaliacaoEstadoSaude is not None:
            participante.autoAvaliacaoEstadoSaude = autoAvaliacaoEstadoSaude

        nrFilhos = new_post.get('nrFilhos', None)
        if nrFilhos is not None:
            participante.nrFilhos = nrFilhos

        dadosCuidador = new_post.get('dadosCuidador', None)
        if dadosCuidador is not None:
            participante.dadosCuidador = dadosCuidador



        doenca = new_post.get('doenca', None)
        if doenca is not None:
            for id_doenca in doenca:
                d = Doenca.objects.get(pk=id_doenca)
                if d not in participante.diagnosticos.all():
                    participante.diagnosticos.add(d)

        outra_doenca = new_post.get('outra_doenca', None)
        if outra_doenca is not None:
            Doenca.objects.create(nome=outra_doenca)
            participante.doenca = outra_doenca
        

        participante.save()
        return HttpResponseRedirect(reverse('protocolo:participant', args=(participante_id,)))

    contexto = {
        'doencas': Doenca.objects.all(),
        'formParticipante': formParticipante,
        'participante': participante,
        'dadosCuidador': participante.dadosCuidador,
    }
    print('contexto', contexto)

    return render(request, "protocolo/participante_update.html", contexto)


@login_required(login_url='diario:login')
def cuidador_update(request, cuidador_id):
    cuidador = Cuidador.objects.get(pk=cuidador_id)
    formCuidador = CuidadorForm(request.POST or None, instance=cuidador)

    if request.method == 'POST':
        new_post = remove_empty_keys_from_post(request.POST)

        nome = new_post.get('nome', None)
        if nome is not None:
            cuidador.info_sensivel.nome = nome

        email = new_post.get('email', None)
        if email is not None:
            cuidador.info_sensivel.email = email

        telemovel = new_post.get('telemovel', None)
        if telemovel is not None:
            cuidador.info_sensivel.telemovel = telemovel

        imagem = request.FILES.get('imagem', None)
        if imagem is not None:
            cuidador.imagem = imagem

        escolaridade = new_post.get('escolaridade', None)
        if escolaridade is not None:
            cuidador.escolaridade = escolaridade

        nascimento = new_post.get('nascimento', None)
        if nascimento is not None:
            cuidador.nascimento = nascimento

        nacionalidade = new_post.get('nacionalidade', None)
        if nacionalidade is not None:
            cuidador.nacionalidade = nacionalidade

        sexo = new_post.get('sexo', None)
        if sexo is not None:
            cuidador.sexo = sexo

        localizacao = new_post.get('localizacao', None)
        if localizacao is not None:
            cuidador.localizacao = localizacao

        situacaoLaboral = new_post.get('situacaoLaboral', None)
        if situacaoLaboral is not None:
            cuidador.situacaoLaboral = situacaoLaboral

        dadosCuidador = new_post.get('dadosCuidador', None)
        if dadosCuidador is not None:
            cuidador.dadosCuidador = dadosCuidador

        profissaoPrincipal = new_post.get('profissaoPrincipal', None)
        if profissaoPrincipal is not None:
            cuidador.profissaoPrincipal = profissaoPrincipal

        situacaoEconomica = new_post.get('situacaoEconomica', None)
        if situacaoEconomica is not None:
            cuidador.situacaoEconomica = situacaoEconomica

        estadoCivil = new_post.get('estadoCivil', None)
        if estadoCivil is not None:
            cuidador.estadoCivil = estadoCivil

        agregadoFamiliar = new_post.get('agregadoFamiliar', None)
        if agregadoFamiliar is not None:
            cuidador.agregadoFamiliar = agregadoFamiliar

        autoAvaliacaoEstadoSaude = new_post.get(
            'autoAvaliacaoEstadoSaude', None)
        if autoAvaliacaoEstadoSaude is not None:
            cuidador.autoAvaliacaoEstadoSaude = autoAvaliacaoEstadoSaude

        doenca = new_post.get('doenca', None)
        if doenca is not None:
            for id_doenca in doenca:
                d = Doenca.objects.get(pk=id_doenca)
                if d not in cuidador.diagnosticos.all():
                    cuidador.diagnosticos.add(d)

        outra_doenca = new_post.get('outra_doenca', None)
        if outra_doenca is not None:
            Doenca.objects.create(nome=outra_doenca)
            cuidador.doenca = outra_doenca

        cuidador.save()
        # URL DO PERFIL DO CUIDADOR
        return HttpResponseRedirect(reverse('protocolo:profile_cuidador', args=(cuidador_id,)))

    contexto = {
        'doencas': Doenca.objects.all(),
        'formCuidador': formCuidador,
        'cuidador': cuidador,
    }
    print('contexto', contexto)

    return render(request, "protocolo/cuidador_update.html", contexto)


@login_required(login_url='login')
def colaboradores_view(request):
    doctor = request.user
    participants = []
    avaliadores = []
    lista_nr_participantes = []

    try:
        user = Avaliador.objects.get(user=doctor)
        participants = Participante.objects.filter(avaliador=doctor)

    except Avaliador.DoesNotExist:
        try:
            user = Administrador.objects.get(user=doctor)
            participants = Participante.objects.filter(
                referenciacao=user.reference)
            all_avaliadores = DjangoGroup.objects.get(
                name="Avaliador").user_set.all()

            for avaliador in all_avaliadores:
                if Avaliador.objects.filter(user=avaliador, reference=user.reference).exists():
                    avaliadores.append(avaliador)
                    lista_nr_participantes.append(get_nr_participantes(doctor))
        except Administrador.DoesNotExist:
            user = None

    resolutions = Resolution.objects.filter(doctor=doctor)
    if request.user.is_superuser:
        participants = Participante.objects.all()
        resolutions = Resolution.objects.all()

    context = {'participants': participants,
               'resolutions': resolutions,
               'avaliadores': avaliadores,
               'avaliadores_zip': zip(avaliadores, lista_nr_participantes), }
    return render(request, 'protocolo/colaboradores.html', context)


@login_required(login_url='login')
def profile_cuidador_view(request, cuidador_id):
    # Falta mostrar as resoluções das partes feitas no perfil e os seus relatorios
    cuidador = Cuidador.objects.filter(pk=cuidador_id).get()
    resolutions = Resolution.objects.filter(
        cuidador=cuidador, doctor=request.user)
    partesDoUtilizador = ParteDoUtilizador.objects.filter(
        cuidador=cuidador, resolution__doctor=request.user)
    parte = Part.objects.all()

    # r = Resolution.objects.filter(patient = patient, part__part__name = "MentHA-Risk")
    existing_risk = ParteDoUtilizador.objects.filter(
        cuidador=cuidador, part__name="MentHA-Risk")

    user = request.user
    # print(request.user.groups.all())

    # isto server para ver qual dos grupos dos users para ver permições e deixar aceder o que ao que
    user_risk = None
    user_tudo = None

    if request.user.groups.filter(name='Administrador').exists():
        user_tudo = request.user.groups.filter(name='Administrador')
    if request.user.groups.filter(name='Avaliador').exists():
        user_tudo = request.user.groups.filter(name='Avaliador')
    if request.user.groups.filter(name='Avaliador-Risk').exists():
        user_risk = request.user.groups.filter(name='Avaliador-Risk')

    risk_area = Area.objects.get(id=47)
    pergunta_risk = Question.objects.get(id=189)

    # estes IDS funcionam no servidor mas nao localmente
    area_soc = Area.objects.get(pk=77)
    pergunta_soc = Question.objects.get(pk=256)

    c = []
    age = calculate_age(cuidador.nascimento)
    answered_list = []
    percentage_list = []

    form = AppointmentForm(request.POST or None)

    if request.method == "POST":
        form = AppointmentForm(request.POST)
        print(request.POST)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.cuidador = cuidador
            obj.save()

            resolution = Resolution(
                part=obj, cuidador=cuidador, doctor=request.user)
            resolution.initialize_statistics()
            resolution.save()

    # if request.user.groups.filter(name='Avaliador-Risk').exists():
    #     if not parteDoUtilizador.filter(part= risk_area.part):
    #         parte

    for parteDoUtilizador in partesDoUtilizador:
        resolution = resolutions.filter(
            part=parteDoUtilizador, cuidador=cuidador, doctor=request.user)
        if not resolution:
            answered_list.append(0)
            percentage_list.append(0)
        else:
            s = resolution.get().statistics
            answered_list.append(s.get('total_answered'))
            percentage_list.append(s.get('total_percentage'))
    for participante in Participante.objects.all():
        if participante in cuidador.participantes.all():
            c.append(participante.nome)

    participantes = ", ".join(c)

    # Copia da patient_overview_view
    areas = Area.objects.order_by('part__order', 'order').all()
    # print(areas)
    overview_list = []

    # Getting all unique area names to display on html page
    for area in areas:
        for instrument in Instrument.objects.filter(area=area):
            text = area.name
            if instrument.name != "None":
                text = text + " - " + instrument.name
            overview_list.append(text)

    ow_l = dict.fromkeys(x for x in overview_list)

    percentages = {}
    # Getting % done per area, per part, to display done or not done on HTML
    for part_do_utilizador in partesDoUtilizador:
        percentages[part_do_utilizador.id] = {}
        part = part_do_utilizador.part
        r = resolutions.filter(part=part_do_utilizador)
        if len(r) < 1:
            for text in overview_list:
                this_area = Area.objects.filter(
                    part=part, name=text.split(' - ')[0]).order_by('order')
                if len(this_area) < 1:
                    percentages[part_do_utilizador.id][text] = "does not exist"
                else:
                    percentages[part_do_utilizador.id][text] = "not done"
        else:
            r = r.get()
            for text in overview_list:
                area_text = text.split(' - ')[0]
                this_area = Area.objects.filter(
                    part=part, name=area_text).order_by('order')
                if len(this_area) < 1:
                    percentages[part_do_utilizador.id][text] = "does not exist"
                else:
                    this_area = this_area.get()
                    instruments = Instrument.objects.filter(
                        area=this_area, area__part=part).order_by('order')
                    for instrument in instruments:
                        if instrument.name != 'None':
                            p = r.statistics[f"{this_area.id}"][f"{instrument.id}"].get(
                                'percentage')

                        else:
                            p = r.statistics[f"{this_area.id}"].get(
                                'percentage')

                        if p == 100:
                            percentages[part_do_utilizador.id][text] = "done"
                        else:
                            percentages[part_do_utilizador.id][text] = "not done"

    context = {'cuidador': cuidador, 'participantes': participantes, 'resolutions': resolutions,
               'partesdoutilizador': zip(partesDoUtilizador, answered_list, percentage_list),
               'partedoutilizador': partesDoUtilizador,
               'overview_list': ow_l,
               'percentages': percentages, 'age': age,
               'parte': parte,
               'risk_area': risk_area,
               'pergunta_risk': pergunta_risk,
               'form': form,
               'user': user,
               'user_risk': user_risk,
               'user_tudo': user_tudo,
               'existing_risk': existing_risk,
               'area_soc': area_soc,
               'pergunta_soc': pergunta_soc,
               }
    return render(request, 'protocolo/profile_cuidador.html', context)
